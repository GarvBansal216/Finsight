import os
import sys
import json
import pytesseract
from pdf2image import convert_from_path
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional, List, Dict, Any
import re
import vertexai
from vertexai.generative_models import GenerativeModel
from fpdf import FPDF
from tempfile import NamedTemporaryFile
from dotenv import load_dotenv
import uvicorn
import time
import platform
import pandas as pd
import asyncio
import base64
import hashlib
import redis
from datetime import datetime

# Initialize Redis client for caching
redis_client = None
REDIS_AVAILABLE = False
try:
    redis_url = os.getenv("REDIS_URL", "redis://localhost:6379/0")
    redis_client = redis.from_url(redis_url, decode_responses=True)
    redis_client.ping()  # Test connection
    REDIS_AVAILABLE = True
    print("✓ Redis connected for caching")
except Exception as e:
    print(f"WARNING: Redis not available for caching: {str(e)}")
    print("Caching will be disabled. Install and start Redis for better performance.")
    redis_client = None
    REDIS_AVAILABLE = False

# Import Celery app (lazy import to avoid issues if Redis is not available)
try:
    from celery_app import celery_app
    CELERY_AVAILABLE = True
except Exception as e:
    print(f"WARNING: Celery not available: {str(e)}")
    print("Async job queue will not work. Install Redis and start Celery worker.")
    celery_app = None
    CELERY_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed. DOCX file processing will not work.")

# Detect operating system
IS_WINDOWS = platform.system() == "Windows"
EXE_EXT = ".exe" if IS_WINDOWS else ""
from report_generators import (
    generate_bank_statement_reports,
    generate_gst_return_reports,
    generate_invoice_reports,
    generate_purchase_order_reports,
    generate_salary_slip_reports,
    generate_profit_loss_reports,
    generate_trial_balance_reports,
    generate_balance_sheet_reports,
    generate_audit_papers_reports,
    generate_agreement_contract_reports,
    generate_comprehensive_audit_report,
    generate_gst_reports_from_excel
)

# Import Tally integration
try:
    from tally_integration import export_to_tally_xml, TallyXMLGenerator
    TALLY_AVAILABLE = True
except ImportError:
    TALLY_AVAILABLE = False
    print("WARNING: Tally integration not available. Install required dependencies if needed.")

# Load environment variables from .env file
load_dotenv()

# Configure Tesseract path for Windows (if needed)
tesseract_cmd = os.getenv("TESSERACT_CMD")
if tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

# Configure Poppler path for Windows (if needed)
# Poppler is required for PDF to image conversion
poppler_path = os.getenv("POPPLER_PATH")
if poppler_path:
    # Normalize and convert to absolute path
    poppler_path = os.path.abspath(os.path.normpath(poppler_path))
    poppler_bin = os.path.join(poppler_path, "bin")
    # Add poppler/bin to PATH if not already there
    if os.path.exists(poppler_bin):
        # Add to the beginning of PATH so it's found first
        current_path = os.environ.get("PATH", "")
        if poppler_bin not in current_path:
            os.environ["PATH"] = os.pathsep.join([poppler_bin, current_path])

# Load Vertex AI configuration from environment variables
# Note: Vertex AI credentials are required for document processing
vertexai_project = os.getenv("VERTEXAI_PROJECT_ID")
vertexai_location = os.getenv("VERTEXAI_LOCATION", "us-central1")
gemini_api_key = os.getenv("GEMINI_API_KEY")  # Fallback to Gemini API key if Vertex AI not configured
vertexai_model_name = os.getenv("VERTEXAI_MODEL", "gemini-1.5-pro")

# Initialize Vertex AI client
vertexai_client = None
gemini_api_key_available = False

# Prioritize Gemini API key if available (no need for ADC/Vertex AI setup)
if gemini_api_key:
    gemini_api_key_available = True
    print(f"✓ Using Gemini API directly (API key provided - no ADC setup needed)")
    # Skip Vertex AI initialization when using Gemini API key
    vertexai_client = None
elif vertexai_project and vertexai_project != "your-gcp-project-id":
    # Only use Vertex AI if no Gemini API key is provided
    try:
        vertexai.init(project=vertexai_project, location=vertexai_location)
        vertexai_client = GenerativeModel(vertexai_model_name)
        print(f"✓ Vertex AI initialized: project={vertexai_project}, location={vertexai_location}, model={vertexai_model_name}")
    except Exception as e:
        print(f"WARNING: Vertex AI initialization failed: {str(e)}")
        print("Please set GEMINI_API_KEY in .env file as fallback")
        vertexai_client = None
else:
    print("WARNING: Neither Vertex AI nor Gemini API key configured. Please set GEMINI_API_KEY in .env file.")

# Helper function to generate cache key from prompt
def get_cache_key(prompt: str, require_json: bool = False) -> str:
    """Generate a cache key from prompt content"""
    cache_string = f"{prompt}_{require_json}"
    cache_hash = hashlib.md5(cache_string.encode()).hexdigest()
    return f"gemini_cache:{cache_hash}"

# Helper function to generate document cache key from file content
def get_document_cache_key(file_content: bytes, document_type: Optional[str] = None) -> str:
    """Generate a cache key from file content hash"""
    file_hash = hashlib.sha256(file_content).hexdigest()
    doc_type_str = document_type or "auto"
    return f"document_cache:{file_hash}:{doc_type_str}"

# Helper function to generate content using Vertex AI
def generate_content_with_vertexai(prompt: str, require_json: bool = False, max_retries: int = 3, use_cache: bool = True):
    """
    Generate content using Vertex AI Gemini model with Redis caching.
    
    Args:
        prompt: The prompt to send to the model
        require_json: If True, expects JSON response and adds JSON instruction to prompt
        max_retries: Maximum number of retry attempts
        use_cache: Whether to use Redis cache (default: True)
    
    Returns:
        str: The generated content
    """
    # Check cache first if Redis is available
    if use_cache and REDIS_AVAILABLE and redis_client:
        try:
            cache_key = get_cache_key(prompt, require_json)
            cached_response = redis_client.get(cache_key)
            if cached_response:
                print(f"✓ Cache HIT for prompt (key: {cache_key[:20]}...)")
                return cached_response
        except Exception as e:
            print(f"Cache read error (continuing without cache): {str(e)}")
    
    # Prioritize Gemini API if available (simpler, no ADC setup needed)
    if gemini_api_key_available:
        import google.generativeai as genai
        genai.configure(api_key=gemini_api_key)
        
        # Add JSON instruction if required
        if require_json:
            prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
        
        # Use gemini-2.5-flash as primary (best free model - latest, fast, free tier, high quality)
        # Fallback to gemini-2.0-flash if 2.5 is unavailable
        # These are the actual available models in the free Gemini API
        model_names = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-2.0-flash-lite']
        
        for attempt in range(max_retries):
            for model_name in model_names:
                try:
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content(prompt)
                    
                    if not response or not response.text:
                        raise ValueError("Gemini API returned empty content")
                    
                    content = response.text.strip()
                    
                    # Clean up JSON if needed
                    if require_json:
                        # Remove markdown code blocks if present
                        if content.startswith("```json"):
                            content = content.replace("```json", "").replace("```", "").strip()
                        elif content.startswith("```"):
                            content = content.replace("```", "").strip()
                    
                    # Cache the response if Redis is available
                    if use_cache and REDIS_AVAILABLE and redis_client:
                        try:
                            cache_key = get_cache_key(prompt, require_json)
                            # Cache for 24 hours (86400 seconds)
                            redis_client.setex(cache_key, 86400, content)
                            print(f"✓ Cached Gemini API response (key: {cache_key[:20]}..., TTL: 24h)")
                        except Exception as e:
                            print(f"Cache write error (response still returned): {str(e)}")
                    
                    return content
                except Exception as e:
                    error_str = str(e)
                    # If model not found, try next model name
                    if "404" in error_str and "not found" in error_str.lower():
                        print(f"Model {model_name} not available, trying next model...")
                        continue
                    # For other errors on last attempt, raise
                    if attempt == max_retries - 1 and model_name == model_names[-1]:
                        raise
                    # Otherwise, wait and retry
                    if attempt < max_retries - 1:
                        wait_time = (attempt + 1) * 1
                        time.sleep(wait_time)
                        break
        
        raise HTTPException(
            status_code=500,
            detail="Failed to generate content with Gemini API after trying all available models."
        )
    
    # Add JSON instruction if required
    if require_json:
        prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
    
    for attempt in range(max_retries):
        try:
            response = vertexai_client.generate_content(prompt)
            
            if not response or not response.text:
                raise ValueError("Vertex AI returned empty content")
            
            content = response.text.strip()
            
            # Clean up JSON if needed
            if require_json:
                # Remove markdown code blocks if present
                if content.startswith("```json"):
                    content = content.replace("```json", "").replace("```", "").strip()
                elif content.startswith("```"):
                    content = content.replace("```", "").strip()
            
            # Cache the response if Redis is available
            if use_cache and REDIS_AVAILABLE and redis_client:
                try:
                    cache_key = get_cache_key(prompt, require_json)
                    # Cache for 24 hours (86400 seconds)
                    redis_client.setex(cache_key, 86400, content)
                    print(f"✓ Cached Vertex AI response (key: {cache_key[:20]}..., TTL: 24h)")
                except Exception as e:
                    print(f"Cache write error (response still returned): {str(e)}")
            
            return content
        except Exception as e:
            error_str = str(e)
            # Handle authentication errors - fallback to Gemini API if available
            if "401" in error_str or "403" in error_str or "permission" in error_str.lower() or "authentication" in error_str.lower():
                print(f"⚠️ Vertex AI authentication failed: {error_str}")
                if gemini_api_key_available:
                    print(f"🔄 Falling back to Gemini API...")
                    # Fallback to Gemini API
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=gemini_api_key)
                        
                        # Add JSON instruction if required
                        if require_json:
                            fallback_prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
                        else:
                            fallback_prompt = prompt
                        
                        # Use gemini-2.5-flash as primary (best free model)
                        model_names = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-2.0-flash-lite']
                        
                        for model_name in model_names:
                            try:
                                model = genai.GenerativeModel(model_name)
                                response = model.generate_content(fallback_prompt)
                                
                                if not response or not response.text:
                                    continue
                                
                                content = response.text.strip()
                                
                                # Clean up JSON if needed
                                if require_json:
                                    if content.startswith("```json"):
                                        content = content.replace("```json", "").replace("```", "").strip()
                                    elif content.startswith("```"):
                                        content = content.replace("```", "").strip()
                                
                                # Cache the response if Redis is available
                                if use_cache and REDIS_AVAILABLE and redis_client:
                                    try:
                                        cache_key = get_cache_key(prompt, require_json)
                                        redis_client.setex(cache_key, 86400, content)
                                        print(f"✓ Cached Gemini API fallback response (key: {cache_key[:20]}..., TTL: 24h)")
                                    except Exception as cache_err:
                                        print(f"Cache write error (response still returned): {str(cache_err)}")
                                
                                print(f"✓ Successfully used Gemini API fallback (model: {model_name})")
                                return content
                            except Exception as gemini_error:
                                error_str_gemini = str(gemini_error)
                                if "404" in error_str_gemini and "not found" in error_str_gemini.lower():
                                    print(f"Model {model_name} not available, trying next model...")
                                    continue
                                # If last model failed, continue to raise error
                                if model_name == model_names[-1]:
                                    raise
                        
                        raise HTTPException(
                            status_code=500,
                            detail="Failed to generate content with Gemini API fallback after trying all models."
                        )
                    except HTTPException:
                        raise
                    except Exception as fallback_error:
                        print(f"⚠️ Gemini API fallback also failed: {str(fallback_error)}")
                        raise HTTPException(
                            status_code=500,
                            detail=f"Vertex AI authentication failed and Gemini API fallback also failed: {str(fallback_error)}"
                        )
                else:
                    # No Gemini API fallback available
                    raise HTTPException(
                        status_code=401,
                        detail="Vertex AI authentication failed. Please check your GCP credentials and project configuration, or set GEMINI_API_KEY as a fallback."
                    )
            # Handle rate limiting
            elif "429" in error_str or "rate_limit" in error_str.lower() or "quota" in error_str.lower():
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 2  # Exponential backoff
                    print(f"Rate limit hit, waiting {wait_time} seconds before retry...")
                    time.sleep(wait_time)
                    continue
                else:
                    # If rate limited and Gemini API available, try fallback
                    if gemini_api_key_available:
                        print(f"🔄 Vertex AI rate limited, falling back to Gemini API...")
                        try:
                            import google.generativeai as genai
                            genai.configure(api_key=gemini_api_key)
                            
                            if require_json:
                                fallback_prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
                            else:
                                fallback_prompt = prompt
                            
                            model_names = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-2.0-flash-lite']
                            for model_name in model_names:
                                try:
                                    model = genai.GenerativeModel(model_name)
                                    response = model.generate_content(fallback_prompt)
                                    if response and response.text:
                                        content = response.text.strip()
                                        if require_json:
                                            if content.startswith("```json"):
                                                content = content.replace("```json", "").replace("```", "").strip()
                                            elif content.startswith("```"):
                                                content = content.replace("```", "").strip()
                                        print(f"✓ Successfully used Gemini API fallback due to rate limit (model: {model_name})")
                                        return content
                                except:
                                    if model_name == model_names[-1]:
                                        raise
                            raise
                        except:
                            raise HTTPException(
                                status_code=429,
                                detail="Vertex AI API rate limit exceeded and Gemini API fallback failed. Please wait a few seconds and try again."
                            )
                    else:
                        raise HTTPException(
                            status_code=429,
                            detail="Vertex AI API rate limit exceeded. Please wait a few seconds and try again."
                        )
            else:
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 1
                    time.sleep(wait_time)
                    continue
                else:
                    # On final failure, try Gemini API fallback if available
                    if gemini_api_key_available:
                        print(f"🔄 Vertex AI failed after retries, falling back to Gemini API...")
                        try:
                            import google.generativeai as genai
                            genai.configure(api_key=gemini_api_key)
                            
                            if require_json:
                                fallback_prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
                            else:
                                fallback_prompt = prompt
                            
                            model_names = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-2.0-flash-lite']
                            for model_name in model_names:
                                try:
                                    model = genai.GenerativeModel(model_name)
                                    response = model.generate_content(fallback_prompt)
                                    if response and response.text:
                                        content = response.text.strip()
                                        if require_json:
                                            if content.startswith("```json"):
                                                content = content.replace("```json", "").replace("```", "").strip()
                                            elif content.startswith("```"):
                                                content = content.replace("```", "").strip()
                                        print(f"✓ Successfully used Gemini API fallback (model: {model_name})")
                                        return content
                                except:
                                    if model_name == model_names[-1]:
                                        raise
                            raise
                        except Exception as fallback_error:
                            raise HTTPException(
                                status_code=500,
                                detail=f"Error generating content with Vertex AI: {str(e)}. Gemini API fallback also failed: {str(fallback_error)}"
                            )
                    else:
                        raise HTTPException(
                            status_code=500,
                            detail=f"Error generating content with Vertex AI: {str(e)}"
                        )
    
    # If all Vertex AI retries failed, try Gemini API fallback
    if gemini_api_key_available:
        print(f"🔄 Vertex AI failed after all retries, falling back to Gemini API...")
        try:
            import google.generativeai as genai
            genai.configure(api_key=gemini_api_key)
            
            if require_json:
                fallback_prompt = f"{prompt}\n\nIMPORTANT: Return ONLY valid JSON. Do not include markdown formatting, code blocks, or any explanations."
            else:
                fallback_prompt = prompt
            
            model_names = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-2.0-flash-lite']
            for model_name in model_names:
                try:
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content(fallback_prompt)
                    if response and response.text:
                        content = response.text.strip()
                        if require_json:
                            if content.startswith("```json"):
                                content = content.replace("```json", "").replace("```", "").strip()
                            elif content.startswith("```"):
                                content = content.replace("```", "").strip()
                        print(f"✓ Successfully used Gemini API fallback after Vertex AI retries (model: {model_name})")
                        return content
                except:
                    if model_name == model_names[-1]:
                        raise
            raise
        except Exception as fallback_error:
            raise HTTPException(status_code=500, detail=f"Failed to get response from Vertex AI after retries. Gemini API fallback also failed: {str(fallback_error)}")
    
    raise HTTPException(status_code=500, detail="Failed to get response from Vertex AI after retries")

app = FastAPI(title="FinSight Document Processor", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this properly for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add logging middleware to see all requests
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@app.middleware("http")
async def log_requests(request, call_next):
    import sys
    sys.stdout.flush()  # Force flush
    print(f"\n{'='*60}", flush=True)
    print(f"REQUEST: {request.method} {request.url.path}", flush=True)
    print(f"Client: {request.client.host if request.client else 'unknown'}", flush=True)
    print(f"{'='*60}\n", flush=True)
    try:
        response = await call_next(request)
        print(f"RESPONSE: {response.status_code} for {request.method} {request.url.path}", flush=True)
        sys.stdout.flush()
        return response
    except Exception as e:
        print(f"ERROR in request: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        sys.stdout.flush()
        raise

# ------------------------------
# OCR HELPERS
# ------------------------------

def extract_text_from_pdf(file_path):
    """Extract text from PDF using OCR"""
    try:
        # Try to get poppler path from environment
        poppler_path = os.getenv("POPPLER_PATH")
        
        # Verify poppler path exists if specified
        if poppler_path:
            # Normalize and get absolute path (handle spaces, etc.)
            poppler_path = os.path.abspath(os.path.normpath(poppler_path))
            poppler_bin = os.path.join(poppler_path, "bin")
            
            if not os.path.exists(poppler_bin):
                raise HTTPException(
                    status_code=500,
                    detail=f"Poppler bin directory not found at: {poppler_bin}. Please check your POPPLER_PATH in .env file."
                )
            
            # Verify both pdfinfo and pdftoppm exist (with appropriate extension for OS)
            pdfinfo_exe = os.path.join(poppler_bin, f"pdfinfo{EXE_EXT}")
            pdftoppm_exe = os.path.join(poppler_bin, f"pdftoppm{EXE_EXT}")
            if not os.path.exists(pdfinfo_exe) or not os.path.exists(pdftoppm_exe):
                raise HTTPException(
                    status_code=500,
                    detail=f"Poppler executables not found. pdfinfo{EXE_EXT}: {os.path.exists(pdfinfo_exe)}, pdftoppm{EXE_EXT}: {os.path.exists(pdftoppm_exe)}. Please verify Poppler installation."
                )
            
            # Ensure PATH includes poppler/bin (update it right before use)
            current_path = os.environ.get("PATH", "")
            if poppler_bin not in current_path:
                os.environ["PATH"] = os.pathsep.join([poppler_bin, current_path])
            
            # Try using poppler_path parameter first (most reliable on Windows with spaces)
            try:
                pages = convert_from_path(file_path, poppler_path=poppler_path)
            except Exception as e1:
                # Fallback: Try using PATH (in case poppler_path parameter doesn't work)
                try:
                    pages = convert_from_path(file_path)
                except Exception as e2:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to use Poppler. Direct path error: {str(e1)}. PATH method error: {str(e2)}. Poppler bin: {poppler_bin}. Please verify the path is correct and restart the server."
                    )
        else:
            # Try default (assumes poppler is in PATH)
            pages = convert_from_path(file_path)
        
        text = ""
        tesseract_errors = []
        for i, page in enumerate(pages):
            try:
                # Verify Tesseract is configured
                if not pytesseract.pytesseract.tesseract_cmd:
                    tesseract_cmd = os.getenv("TESSERACT_CMD")
                    if tesseract_cmd:
                        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                    else:
                        raise HTTPException(
                            status_code=500,
                            detail="Tesseract OCR is not configured. Please set TESSERACT_CMD in your .env file."
                        )
                
                page_text = pytesseract.image_to_string(page)
                if page_text.strip():
                    text += page_text
                else:
                    print(f"Warning: Page {i+1} returned empty text from OCR")
            except HTTPException:
                raise
            except Exception as tesseract_error:
                error_msg = str(tesseract_error)
                tesseract_errors.append(f"Page {i+1}: {error_msg}")
                print(f"Warning: Tesseract OCR failed for page {i+1}: {error_msg}")
                # Continue processing other pages
        
        if not text.strip():
            error_detail = "No text could be extracted from the PDF."
            if tesseract_errors:
                error_detail += f" Tesseract errors: {'; '.join(tesseract_errors[:3])}"
            else:
                error_detail += " Please check if Tesseract OCR is properly installed and configured."
            error_detail += f" Tesseract path: {pytesseract.pytesseract.tesseract_cmd or 'Not set'}"
            raise HTTPException(
                status_code=500,
                detail=error_detail
            )
        
        return text
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        # Check for elevation/permission errors
        if "elevation" in error_msg.lower() or "WinError 740" in error_msg or "requires elevation" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail=f"Permission error: {error_msg}. This usually means the process needs administrator privileges. Try:\n1. Running the FastAPI server as Administrator (right-click → Run as Administrator)\n2. Or check if antivirus/Windows Defender is blocking the executables\n3. Ensure temporary files can be created in your user directory"
            )
        if "poppler" in error_msg.lower() or "Unable to get page count" in error_msg:
            poppler_path = os.getenv("POPPLER_PATH")
            if poppler_path:
                detail_msg = f"Poppler configuration issue. Path set: {poppler_path}. Error: {error_msg}. Please ensure:\n1. The path is correct and contains 'bin' folder\n2. Both pdfinfo{EXE_EXT} and pdftoppm{EXE_EXT} exist in the bin folder\n3. You restarted the server after updating .env\n4. If the path contains spaces, try moving Poppler to a path without spaces"
            else:
                detail_msg = "Poppler is not installed or not in PATH. Please install Poppler and add it to your PATH, or set POPPLER_PATH in .env file. See README_PYTHON.md for installation instructions."
            raise HTTPException(status_code=500, detail=detail_msg)
        raise HTTPException(status_code=500, detail=f"Error extracting text from PDF: {str(e)}")

def extract_data_from_excel(file_path):
    """Extract data from Excel file and return as structured data"""
    try:
        # Read all sheets from the Excel file
        excel_data = {}
        xls = pd.ExcelFile(file_path)
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            # Convert DataFrame to dictionary format
            excel_data[sheet_name] = {
                "columns": df.columns.tolist(),
                "data": df.fillna("").to_dict('records'),  # Fill NaN with empty string
                "row_count": len(df),
                "column_count": len(df.columns)
            }
        
        # Also create a comprehensive text representation for AI processing
        summary_text = f"Excel file with {len(xls.sheet_names)} sheet(s):\n"
        for sheet_name, sheet_data in excel_data.items():
            summary_text += f"\nSheet: {sheet_name}\n"
            summary_text += f"Columns: {', '.join(sheet_data['columns'])}\n"
            summary_text += f"Total Rows: {sheet_data['row_count']}\n"
            # Include ALL rows (not just sample) - this is critical for accurate data extraction
            if sheet_data['data']:
                summary_text += f"\nAll Data ({sheet_data['row_count']} rows):\n"
                for i, row in enumerate(sheet_data['data']):
                    # Format row data more readably
                    row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if val != "" and val is not None])
                    summary_text += f"  Row {i+1}: {row_str}\n"
        
        return {
            "excel_data": excel_data,
            "summary_text": summary_text,
            "sheet_names": xls.sheet_names
        }
    except Exception as e:
        print(f"Error extracting data from Excel: {str(e)}")
        return {
            "excel_data": {},
            "summary_text": f"Error reading Excel file: {str(e)}",
            "sheet_names": [],
            "error": str(e)
        }

def extract_text_from_image(file_path):
    """Extract text from image using OCR"""
    try:
        return pytesseract.image_to_string(file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text from image: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        if not DOCX_AVAILABLE:
            raise HTTPException(
                status_code=500,
                detail="python-docx library is not installed. Please install it with: pip install python-docx"
            )
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            raise HTTPException(
                status_code=500,
                detail="No text could be extracted from the DOCX file."
            )
        
        return text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting text from DOCX: {str(e)}")

# ------------------------------
# DOCUMENT TYPE CLASSIFIER
# ------------------------------

def classify_document(text):
    """Classify document type using Vertex AI"""
    prompt = f"""
    You are a financial document classifier.
    Classify the document as one of:
    - Bank Statement
    - Invoice
    - GST Document (GSTR-1/GSTR-3B)
    - Trial Balance
    - Profit & Loss Statement
    - Purchase Order
    - Salary Slip (Payroll Report)
    - Balance Sheet
    - Contract/Agreement
    - Unknown
    
    Return JSON only:
    {{
       "type": ""
    }}
    
    Document:
    {text[:3000]}
    """
    
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        return json.loads(response_text)
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in classify_document: {str(e)}")
        print(f"Full traceback:\n{error_trace}")
        raise HTTPException(status_code=500, detail=f"Error classifying document: {str(e)}. Check backend logs for details.")

# ------------------------------
# EXTRACT BANK STATEMENT TABLE
# ------------------------------

def extract_bank_statement(text):
    """Extract bank statement data using Vertex AI"""
    prompt = f"""You are an expert in financial document reconstruction.

You will receive OCR text extracted from a bank statement. This OCR text may contain
symbols like ■, ●, |, -, strange spacing, or broken Indian currency formatting such as
"■ 2,50,000" or "■1,20,000".

Your job:
1. Identify TRUE debit and credit values even if OCR symbols appear before them.
2. Convert Indian number format (2,50,000) into plain integer (250000).
3. If amounts are missing but can be inferred from the balance change, RECONSTRUCT them.
4. Do NOT output 0 unless the OCR truly shows zero.
5. Never guess rows that do not exist.

Output format - MUST be a valid JSON object with two keys:
1. "transactions": array of transaction objects
2. "summary": object with balances and totals

Transaction format:
[
  {{
    "date": "YYYY-MM-DD",
    "description": "",
    "type": "credit" or "debit",
    "amount": number,
    "balance": number or null
  }}
]

Rules:
- For credit rows, amount is the credited value.
- For debit rows, amount is the debited value.
- Remove commas, rupee symbols, bullets (■), pipes (|), and formatting artifacts.
- Balance must be parsed into clean integers.
- If OCR shows a bullet (■), treat it as a formatting symbol, NOT a number.
- Date format must ALWAYS be YYYY-MM-DD.
- If balance is shown, parse it; otherwise set `null`.

Summary format:
{{
  "opening_balance": number or null,
  "closing_balance": number or null,
  "total_debits": number,
  "total_credits": number,
  "anomalies": []
}}

Where anomalies include:
- balance mismatch (balance != prev_balance + credit - debit)
- missing dates
- corrupted rows
- reconstructed amounts (if you inferred them from balance changes)

DO NOT output anything except the JSON object with "transactions" and "summary" keys.

Now extract the transactions from the following OCR text:
{text}
"""
    
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        
        # Check if response is empty
        if not response_text:
            raise HTTPException(
                status_code=500,
                detail="Vertex AI returned an empty response. Please try again."
            )
        
        try:
            parsed_data = json.loads(response_text)
            
            # Handle new format: {transactions: [...], summary: {...}}
            if isinstance(parsed_data, dict) and "transactions" in parsed_data:
                # Store both transactions and summary for later use
                transactions = parsed_data.get("transactions", [])
                summary = parsed_data.get("summary", {})
                
                # Convert to old format for compatibility: debit/credit fields
                formatted_transactions = []
                for txn in transactions:
                    formatted_txn = {
                        "date": txn.get("date", ""),
                        "description": txn.get("description", ""),
                        "debit": txn.get("amount", 0) if txn.get("type") == "debit" else None,
                        "credit": txn.get("amount", 0) if txn.get("type") == "credit" else None,
                        "balance": txn.get("balance"),
                        "category": ""  # Will be categorized later if needed
                    }
                    formatted_transactions.append(formatted_txn)
                
                # Return dict with both transactions and summary
                return {
                    "_transactions": formatted_transactions,
                    "_summary": summary
                }
            else:
                # Old format: just array of transactions
                return parsed_data
                
        except json.JSONDecodeError as json_error:
            # Try to extract JSON from the response using regex
            import re
            # Try to find JSON object first
            json_match = re.search(r'\{.*"transactions".*\}', response_text, re.DOTALL)
            if not json_match:
                # Fallback to array format
                json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            
            if json_match:
                try:
                    parsed_data = json.loads(json_match.group())
                    if isinstance(parsed_data, dict) and "transactions" in parsed_data:
                        transactions = parsed_data.get("transactions", [])
                        summary = parsed_data.get("summary", {})
                        formatted_transactions = []
                        for txn in transactions:
                            formatted_txn = {
                                "date": txn.get("date", ""),
                                "description": txn.get("description", ""),
                                "debit": txn.get("amount", 0) if txn.get("type") == "debit" else None,
                                "credit": txn.get("amount", 0) if txn.get("type") == "credit" else None,
                                "balance": txn.get("balance"),
                                "category": ""
                            }
                            formatted_transactions.append(formatted_txn)
                        return {
                            "_transactions": formatted_transactions,
                            "_summary": summary
                        }
                    return parsed_data
                except:
                    pass
            
            print(f"JSON parsing error in extract_bank_statement. Response: {response_text[:500]}")
            raise HTTPException(
                status_code=500,
                detail=f"Error parsing JSON from Vertex AI response: {str(json_error)}. Response preview: {response_text[:200]}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        error_str = str(e)
        if "quota" in error_str.lower() or "insufficient_quota" in error_str.lower() or "429" in error_str:
            raise HTTPException(
                status_code=402,
                detail="Vertex AI API quota exceeded. Please check your GCP billing and quotas."
            )
        raise HTTPException(status_code=500, detail=f"Error extracting bank statement: {str(e)}")

# ------------------------------
# VALIDATION + RECONCILIATION
# ------------------------------

def validate_statement(rows):
    """Validate and reconcile bank statement - uses summary from extraction if available"""
    # Check if rows have summary data from extraction (new format)
    if isinstance(rows, dict) and "_summary" in rows and "_transactions" in rows:
        summary = rows["_summary"]
        transactions = rows["_transactions"]
        
        # Calculate totals from transactions
        total_debits = sum(t.get("debit", 0) or 0 for t in transactions)
        total_credits = sum(t.get("credit", 0) or 0 for t in transactions)
        
        # Use anomalies from extraction, validate balances
        anomalies = summary.get("anomalies", [])
        
        # Additional validation: check balance consistency
        prev_balance = summary.get("opening_balance")
        for txn in transactions:
            if prev_balance is not None and txn.get("balance") is not None:
                expected_balance = prev_balance + (txn.get("credit", 0) or 0) - (txn.get("debit", 0) or 0)
                actual_balance = txn.get("balance")
                if abs(expected_balance - actual_balance) > 1:  # Allow 1 INR tolerance
                    anomalies.append(f"Balance mismatch at {txn.get('date')}: expected {expected_balance}, got {actual_balance}")
            prev_balance = txn.get("balance")
        
        return {
            "validated_rows": transactions,
            "anomalies": anomalies,
            "summary": {
                "opening_balance": summary.get("opening_balance"),
                "closing_balance": summary.get("closing_balance"),
                "total_debits": total_debits,
                "total_credits": total_credits
            }
        }
    
    # Fallback: validate manually if summary not available (old format)
    prompt = f"""
    Validate the following bank statement transactions:
    {json.dumps(rows, indent=2)}
    
    Steps:
    1. Check if balance = previous_balance + credit - debit (allow 1 INR tolerance)
    2. Detect anomalies (difference > 1 INR, missing dates, corrupted rows)
    3. Calculate totals (debits, credits)
    4. Identify opening and closing balances
    
    Return JSON:
    {{
      "validated_rows": [],
      "anomalies": [],
      "summary": {{
          "opening_balance": number or null,
          "closing_balance": number or null,
          "total_debits": number,
          "total_credits": number
      }}
    }}
    """
    
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        
        # Log the response for debugging
        print(f"Vertex AI validation response (first 200 chars): {response_text[:200]}")
        
        # Check if response is empty
        if not response_text:
            raise HTTPException(
                status_code=500,
                detail="Vertex AI returned an empty response. Please try again."
            )
        
        # Try to parse JSON from the response
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()
        
        # Try to parse JSON
        try:
            parsed_json = json.loads(response_text)
            return parsed_json
        except json.JSONDecodeError as json_error:
            # If JSON parsing fails, try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                try:
                    return json.loads(json_match.group())
                except:
                    pass
            
            # If all else fails, return a default structure
            print(f"JSON parsing error. Response was: {response_text[:500]}")
            raise HTTPException(
                status_code=500,
                detail=f"Error parsing JSON from Vertex AI response. The API may have returned invalid JSON. Response preview: {response_text[:200]}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        error_str = str(e)
        if "quota" in error_str.lower() or "insufficient_quota" in error_str.lower() or "429" in error_str:
            raise HTTPException(
                status_code=402,
                detail="Vertex AI API quota exceeded. Please check your GCP billing and quotas."
            )
        raise HTTPException(status_code=500, detail=f"Error validating statement: {str(e)}")

# ------------------------------
# PDF REPORT GENERATOR
# ------------------------------

def generate_pdf(summary, anomalies):
    """Generate PDF report"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, "Bank Statement Report", ln=1)
        pdf.ln(5)
        pdf.cell(0, 10, f"Opening Balance: {summary.get('opening_balance', 0)}", ln=1)
        pdf.cell(0, 10, f"Closing Balance: {summary.get('closing_balance', 0)}", ln=1)
        pdf.cell(0, 10, f"Total Debits: {summary.get('total_debits', 0)}", ln=1)
        pdf.cell(0, 10, f"Total Credits: {summary.get('total_credits', 0)}", ln=1)
        pdf.ln(5)
        pdf.cell(0, 10, "Anomalies Detected:", ln=1)
        pdf.set_font("Arial", size=10)
        for a in anomalies:
            pdf.multi_cell(0, 8, str(a))
        
        output_path = "report.pdf"
        pdf.output(output_path)
        return output_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

# ------------------------------
# DOCUMENT EXTRACTION FUNCTIONS
# ------------------------------

def generate_report_with_ai(text, report_type, document_type, base_data=None):
    """Generate specific reports using AI based on extracted data"""
    report_prompts = {
        "bank_statement": {
            "cash_flow_statement": """Generate a Cash Flow Statement from the bank statement data. Include:
- Operating Activities (inflows and outflows)
- Investing Activities
- Financing Activities
- Net Cash Flow
- Opening and Closing Cash Balance
Return as structured JSON with sections and line items.""",
            "ledger_entries": """Generate Ledger Entries from bank transactions. Format as:
- Date, Description, Debit, Credit, Balance
- Group by account/category
- Include running balance
Return as JSON array of ledger entries.""",
            "payment_receipt_summary": """Generate Payment & Receipt Summary:
- Total Receipts (credits)
- Total Payments (debits)
- Category-wise breakdown
- Top 10 payments and receipts
Return as JSON with summary totals and detailed list.""",
            "anomaly_report": """Generate Anomaly & Suspicious Transaction Report:
- Unusual transaction patterns
- High-value transactions
- Duplicate transactions
- Suspicious descriptions
- Recommendations
Return as JSON with anomalies list and risk assessment.""",
            "bank_reconciliation": """Generate Bank Reconciliation Sheet:
- Book Balance
- Add: Deposits in transit
- Less: Outstanding checks
- Bank Charges/Interest
- Adjusted Bank Balance
- Differences and adjustments
Return as JSON with reconciliation items."""
        },
        "gst_return": {
            "gst_reconciliation": """Generate GST Reconciliation Sheet:
- GSTR-1 vs GSTR-3B reconciliation
- Sales reconciliation
- Tax reconciliation
- Differences and reasons
- Recommendations
Return as JSON with reconciliation details.""",
            "itc_utilization": """Generate ITC Utilization Report:
- Total Input Tax Credit available
- ITC utilized
- ITC pending
- ITC reversal (if any)
- Utilization percentage
- Month-wise breakdown
Return as JSON with ITC details.""",
            "output_input_tax_summary": """Generate Output vs Input Tax Summary:
- Total Output Tax
- Total Input Tax Credit
- Net Tax Payable
- Tax components (CGST, SGST, IGST)
- Comparison chart data
Return as JSON with tax summary.""",
            "gst_liability_statement": """Generate GST Liability Statement:
- Tax liability period-wise
- Payment schedule
- Due dates
- Penalties/Interest
- Payment status
Return as JSON with liability details.""",
            "gstr_invoice_match": """Generate GSTR vs Invoice Match Report:
- Matched invoices
- Unmatched invoices
- Mismatch reasons
- Recommendations
Return as JSON with matching details."""
        }
    }
    
    prompt = report_prompts.get(document_type, {}).get(report_type, "")
    if not prompt:
        return None
    
    full_prompt = f"""{prompt}

Base Data:
{json.dumps(base_data, indent=2) if base_data else "None"}

OCR Text (if needed):
{text[:4000]}

Return ONLY valid JSON. No explanations."""
    
    try:
        response_text = generate_content_with_vertexai(full_prompt, require_json=True)
        return json.loads(response_text)
    except Exception as e:
        return {"error": f"Error generating {report_type}: {str(e)}"}

def clean_number(value):
    """Clean and normalize number values"""
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    # Remove currency symbols, commas, spaces
    cleaned = str(value).replace("₹", "").replace("$", "").replace(",", "").replace(" ", "").strip()
    # Remove Indian number format (2,50,000 -> 250000)
    parts = cleaned.split(".")
    if len(parts) > 1:
        cleaned = parts[0].replace(",", "") + "." + parts[1]
    else:
        cleaned = cleaned.replace(",", "")
    try:
        return float(cleaned)
    except:
        return None

def extract_bank_statement_structured(text):
    """Extract bank statement data in structured format with person-relevant details"""
    prompt = f"""You are a financial document extraction expert. Extract bank statement data from the following OCR text.

Extract ONLY the following fields in STRICT JSON format:
{{
  "account_holder": "",
  "account_number": "",
  "bank_name": "",
  "branch_name": "",
  "ifsc_code": "",
  "statement_start_date": "",
  "statement_end_date": "",
  "opening_balance": "",
  "closing_balance": "",
  "total_deposits": "",
  "total_withdrawals": "",
  "net_balance_change": "",
  "average_monthly_balance": "",
  "transactions": [
    {{
      "date": "",
      "description": "",
      "type": "credit/debit",
      "amount": "",
      "balance": "",
      "category": "",
      "reference_number": "",
      "mode": ""
    }}
  ],
  "summary": {{
    "total_transactions": "",
    "largest_credit": "",
    "largest_debit": "",
    "recurring_payments": [],
    "cash_withdrawals": "",
    "online_transfers": ""
  }}
}}

Rules:
- Clean all numbers: remove currency symbols, commas, convert to numbers (e.g., "₹2,50,000" -> 250000)
- Dates must be in YYYY-MM-DD format if possible
- Categorize transactions: Salary, Rent, Utilities, Shopping, Food, Travel, Investment, Loan, etc.
- Identify recurring payments (same amount/description monthly)
- Extract reference numbers, cheque numbers, UPI references
- If data is missing, use null or empty string
- Return ONLY valid JSON, no explanations

OCR Text:
{text[:8000]}
"""
    
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        
        # Check if response is empty or None
        if not response_text or not response_text.strip():
            raise ValueError("Vertex AI returned empty response")
        
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError as json_error:
            # Try to extract JSON from markdown code blocks
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                try:
                    result = json.loads(json_match.group())
                except:
                    raise ValueError(f"Failed to parse JSON response. Error: {str(json_error)}. Response preview: {response_text[:500]}")
            else:
                raise ValueError(f"Failed to parse JSON response. Error: {str(json_error)}. Response preview: {response_text[:500]}")
        
        # Clean numbers
        for key in ["opening_balance", "closing_balance", "total_deposits", "total_withdrawals", "net_balance_change", "average_monthly_balance"]:
            if key in result:
                result[key] = clean_number(result[key])
        if "transactions" in result:
            for txn in result["transactions"]:
                if "amount" in txn:
                    txn["amount"] = clean_number(txn["amount"])
                if "balance" in txn:
                    txn["balance"] = clean_number(txn["balance"])
        if "summary" in result:
            summary = result["summary"]
            for key in ["largest_credit", "largest_debit", "cash_withdrawals", "online_transfers"]:
                if key in summary:
                    summary[key] = clean_number(summary[key])
        
        # Compute dashboard metrics and ratios
        total_deposits = result.get("total_deposits", 0) or 0
        total_withdrawals = result.get("total_withdrawals", 0) or 0
        transactions = result.get("transactions", [])
        
        # Calculate category summary for pie chart
        category_summary = {}
        for txn in transactions:
            if txn.get("type") == "debit" and txn.get("category"):
                cat = txn["category"]
                category_summary[cat] = category_summary.get(cat, 0) + (txn.get("amount", 0) or 0)
        
        # Calculate monthly breakdown
        monthly_data = {}
        for txn in transactions:
            if txn.get("date"):
                month = txn["date"][:7] if len(txn["date"]) >= 7 else "unknown"
                if month not in monthly_data:
                    monthly_data[month] = {"income": 0, "expenses": 0}
                if txn.get("type") == "credit":
                    monthly_data[month]["income"] += (txn.get("amount", 0) or 0)
                else:
                    monthly_data[month]["expenses"] += (txn.get("amount", 0) or 0)
        
        # Calculate ratios
        savings_rate = ((total_deposits - total_withdrawals) / total_deposits * 100) if total_deposits > 0 else 0
        expense_ratio = (total_withdrawals / total_deposits * 100) if total_deposits > 0 else 0
        deposit_withdrawal_ratio = (total_deposits / total_withdrawals) if total_withdrawals > 0 else 0
        
        # Add dashboard-ready structure
        # Generate all required reports for Bank Statement
        try:
            reports = generate_bank_statement_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        # Keep dashboard for backward compatibility
        result["dashboard"] = {
            "tables": {
                "transactions": transactions,
                "category_summary": [{"category": k, "amount": v} for k, v in category_summary.items()],
                "high_value_transactions": sorted(
                    [txn for txn in transactions if (txn.get("amount", 0) or 0) > 10000],
                    key=lambda x: x.get("amount", 0) or 0,
                    reverse=True
                )[:10]
            },
            "charts": {
                "income_expense_trend": [{"month": k, "income": v["income"], "expenses": v["expenses"]} for k, v in monthly_data.items()],
                "cashflow_line": [{"date": txn.get("date"), "balance": txn.get("balance")} for txn in transactions if txn.get("balance")],
                "spend_pie": [{"category": k, "amount": v} for k, v in category_summary.items()]
            },
            "ratios": {
                "savings_rate": round(savings_rate, 2),
                "expense_ratio": round(expense_ratio, 2),
                "deposit_withdrawal_ratio": round(deposit_withdrawal_ratio, 2),
                "average_monthly_balance": result.get("average_monthly_balance"),
                "average_monthly_savings": round((total_deposits - total_withdrawals) / max(len(monthly_data), 1), 2)
            },
            "requires_multiple_pdfs": False
        }
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        error_msg = str(e) if str(e) else repr(e)
        print(f"ERROR in extract_bank_statement_structured: {error_msg}")
        print(f"Traceback: {error_trace}")
        sys.stdout.flush()  # Ensure error is printed immediately
        raise HTTPException(status_code=500, detail=f"Error extracting bank statement: {error_msg}")

def extract_gst_return(text):
    """Extract GST Return data with person-relevant details"""
    prompt = f"""Extract GST Return (GSTR-1/GSTR-3B) data from the following OCR text.

Return STRICT JSON:
{{
  "gstin": "",
  "legal_name": "",
  "trade_name": "",
  "address": "",
  "period": "",
  "filing_date": "",
  "filing_status": "",
  "total_sales": "",
  "taxable_value": "",
  "cgst": "",
  "sgst": "",
  "igst": "",
  "cess": "",
  "output_tax": "",
  "input_tax_credit": "",
  "reverse_charge": "",
  "net_tax_payable": "",
  "interest_payable": "",
  "late_fee": "",
  "penalty": "",
  "total_payable": "",
  "payment_status": "",
  "payment_date": ""
}}

Rules:
- Clean all numbers (remove symbols, commas)
- Dates in YYYY-MM-DD format
- Extract all tax components (CGST, SGST, IGST, CESS)
- Include filing and payment status
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        for key in ["total_sales", "taxable_value", "cgst", "sgst", "igst", "cess", "output_tax", "input_tax_credit", "reverse_charge", "net_tax_payable", "interest_payable", "late_fee", "penalty", "total_payable"]:
            if key in result:
                result[key] = clean_number(result[key])
        
        # Compute dashboard metrics
        output_tax = result.get("output_tax", 0) or 0
        input_tax_credit = result.get("input_tax_credit", 0) or 0
        total_sales = result.get("total_sales", 0) or 0
        taxable_value = result.get("taxable_value", 0) or 0
        
        itc_utilization = (input_tax_credit / output_tax * 100) if output_tax > 0 else 0
        tax_liability_ratio = (result.get("net_tax_payable", 0) or 0) / taxable_value * 100 if taxable_value > 0 else 0
        
        result["dashboard"] = {
            "tables": {
                "hsn_summary": [],
                "sales_table": [{"period": result.get("period"), "total_sales": total_sales, "taxable_value": taxable_value}]
            },
            "charts": {
                "sales_trend": [{"period": result.get("period"), "sales": total_sales}],
                "output_vs_input_tax": {
                    "output_tax": output_tax,
                    "input_tax_credit": input_tax_credit,
                    "net_payable": result.get("net_tax_payable", 0) or 0
                }
            },
            "ratios": {
                "itc_utilization": round(itc_utilization, 2),
                "tax_liability_ratio": round(tax_liability_ratio, 2),
                "sales_growth": 0  # Requires multiple periods
            },
            "requires_multiple_pdfs": True
        }
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting GST return: {str(e)}")

def extract_trial_balance(text):
    """Extract Trial Balance data with person-relevant details"""
    prompt = f"""Extract Trial Balance data from the following OCR text.

CRITICAL REQUIREMENTS:
1. Extract ONLY values that are explicitly present in the document
2. Do NOT use placeholder values like 0, "N/A", or any default numbers unless the document shows zero
3. If a value is missing in the document, use null (not 0, not empty string)
4. All numbers must be actual extracted values - clean them (remove currency symbols, commas) but do not modify the actual numeric value
5. Calculate totals accurately: sum all debits and credits from the balances array
6. Verify calculations: total_debits and total_credits must match sum of individual account debits/credits
7. Account types must be accurately categorized based on account names and standard accounting principles

Return STRICT JSON:
{{
  "entity_name": "",
  "entity_type": "",
  "period": "",
  "prepared_date": "",
  "total_debits": "",
  "total_credits": "",
  "difference": "",
  "is_balanced": true or false,
  "balances": [
    {{
      "account_code": "",
      "account_name": "",
      "account_type": "Asset" | "Liability" | "Income" | "Expense" | "Equity",
      "debit": number or null,
      "credit": number or null,
      "balance": number or null
    }}
  ],
  "summary": {{
    "asset_accounts": "",
    "liability_accounts": "",
    "income_accounts": "",
    "expense_accounts": ""
  }}
}}

Rules:
- Clean all numbers: remove ₹, commas, but preserve actual numeric value (e.g., "₹2,50,000" -> 250000)
- Categorize accounts accurately: Asset, Liability, Income, Expense, Equity based on account name
- Calculate total_debits = sum of all debit values in balances array
- Calculate total_credits = sum of all credit values in balances array
- Calculate difference = total_debits - total_credits
- Set is_balanced = true only if difference is exactly 0 (or very close to 0 due to rounding)
- For summary: sum asset_accounts from all Asset type accounts, liability_accounts from Liability accounts, etc.
- If a field is not found in document, use null (not 0, not empty string)
- Return ONLY valid JSON, no explanations or additional text

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        for key in ["total_debits", "total_credits", "difference"]:
            if key in result:
                result[key] = clean_number(result[key])
        if "balances" in result:
            for bal in result["balances"]:
                for key in ["debit", "credit", "balance"]:
                    if key in bal:
                        bal[key] = clean_number(bal[key])
        if "summary" in result:
            summary = result["summary"]
            for key in ["asset_accounts", "liability_accounts", "income_accounts", "expense_accounts"]:
                if key in summary:
                    summary[key] = clean_number(summary[key])
        
        # Compute dashboard metrics
        total_debits = result.get("total_debits", 0) or 0
        total_credits = result.get("total_credits", 0) or 0
        balances = result.get("balances", [])
        
        debit_credit_ratio = (total_debits / total_credits) if total_credits > 0 else 0
        
        # Top accounts by value
        top_accounts = sorted(
            balances,
            key=lambda x: max(x.get("debit", 0) or 0, x.get("credit", 0) or 0),
            reverse=True
        )[:10]
        
        # Account type breakdown
        account_type_summary = {}
        for bal in balances:
            acc_type = bal.get("account_type", "Other")
            account_type_summary[acc_type] = account_type_summary.get(acc_type, 0) + max(bal.get("debit", 0) or 0, bal.get("credit", 0) or 0)
        
        result["dashboard"] = {
            "tables": {
                "balances": balances,
                "top_accounts": top_accounts
            },
            "charts": {
                "debit_credit_chart": {
                    "total_debits": total_debits,
                    "total_credits": total_credits
                },
                "account_type_pyramid": [{"type": k, "value": v} for k, v in account_type_summary.items()],
                "ledger_contribution": [{"account_name": b.get("account_name"), "value": max(b.get("debit", 0) or 0, b.get("credit", 0) or 0)} for b in top_accounts]
            },
            "ratios": {
                "debit_credit_ratio": round(debit_credit_ratio, 2),
                "major_account_concentration": round((sum(max(b.get("debit", 0) or 0, b.get("credit", 0) or 0) for b in top_accounts[:5]) / max(total_debits, total_credits, 1)) * 100, 2)
            },
            "requires_multiple_pdfs": False
        }
        
        # Generate all required reports for Trial Balance
        try:
            reports = generate_trial_balance_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except HTTPException:
        # Re-raise HTTPExceptions as-is (they already have proper status codes and messages)
        raise
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error parsing trial balance JSON response: {str(e)}. The AI model may have returned invalid JSON."
        )
    except Exception as e:
        error_msg = str(e)
        # Provide more helpful error messages
        if "Vertex AI is not configured" in error_msg or "VERTEXAI_PROJECT_ID" in error_msg:
            raise HTTPException(
                status_code=500,
                detail="Vertex AI is not configured. Please set VERTEXAI_PROJECT_ID and VERTEXAI_LOCATION in .env file, or set GEMINI_API_KEY as a fallback."
            )
        elif "rate_limit" in error_msg.lower() or "quota" in error_msg.lower() or "429" in error_msg:
            raise HTTPException(
                status_code=429,
                detail="API rate limit exceeded. Please try again in a few moments."
            )
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting trial balance: {error_msg}"
            )

def extract_profit_loss(text):
    """Extract Profit & Loss Statement data with person-relevant details"""
    prompt = f"""Extract Profit & Loss Statement data from the following OCR text.

CRITICAL REQUIREMENTS:
1. Extract ONLY values that are explicitly present in the document
2. Do NOT use placeholder values like 0, "N/A", or any default numbers unless the document shows zero
3. If a value is missing in the document, use null (not 0, not empty string)
4. All numbers must be actual extracted values - clean them (remove currency symbols, commas) but do not modify the actual numeric value
5. Calculate derived values accurately:
   - gross_profit = total_revenue - cost_of_goods_sold
   - operating_profit = gross_profit - operating_expenses
   - profit_before_tax = operating_profit - financial_expenses
   - net_profit = profit_before_tax - tax_expenses
   - gross_margin = (gross_profit / total_revenue) * 100 (only if both values exist)
   - net_margin = (net_profit / total_revenue) * 100 (only if both values exist)
6. Verify that expense breakdown amounts sum to total_expenses
7. Calculate percentages only when you have both numerator and denominator values

Return STRICT JSON:
{{
  "entity_name": "",
  "period": "",
  "prepared_date": "",
  "revenue": {{
    "total_revenue": number or null,
    "sales": number or null,
    "service_income": number or null,
    "other_income": number or null,
    "revenue_breakdown": [
      {{
        "category": "",
        "amount": number or null
      }}
    ]
  }},
  "expenses": {{
    "total_expenses": number or null,
    "cost_of_goods_sold": number or null,
    "operating_expenses": number or null,
    "administrative_expenses": number or null,
    "financial_expenses": number or null,
    "tax_expenses": number or null,
    "expense_breakdown": [
      {{
        "category": "",
        "amount": number or null,
        "percentage": number or null
      }}
    ]
  }},
  "profitability": {{
    "gross_profit": number or null,
    "operating_profit": number or null,
    "profit_before_tax": number or null,
    "net_profit": number or null,
    "gross_margin": number or null,
    "net_margin": number or null
  }}
}}

Rules:
- Clean all numbers: remove ₹, commas, but preserve actual numeric value (e.g., "₹2,50,000" -> 250000)
- Extract values directly from document - do not calculate if values are explicitly shown
- For calculated fields (gross_profit, operating_profit, etc.), use the value from document if shown, otherwise calculate
- Calculate percentages only when you have valid numerator and denominator (not null, not zero)
- Expense breakdown percentages = (individual expense / total_expenses) * 100
- Verify that expense_breakdown amounts sum to total_expenses (within rounding tolerance)
- If a field is not found in document and cannot be calculated from other fields, use null (not 0, not empty string)
- Return ONLY valid JSON, no explanations or additional text

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        # Clean revenue numbers
        if "revenue" in result:
            for key in ["total_revenue", "sales", "service_income", "other_income"]:
                if key in result["revenue"]:
                    result["revenue"][key] = clean_number(result["revenue"][key])
            if "revenue_breakdown" in result["revenue"]:
                for item in result["revenue"]["revenue_breakdown"]:
                    if "amount" in item:
                        item["amount"] = clean_number(item["amount"])
        # Clean expense numbers
        if "expenses" in result:
            for key in ["total_expenses", "cost_of_goods_sold", "operating_expenses", "administrative_expenses", "financial_expenses", "tax_expenses"]:
                if key in result["expenses"]:
                    result["expenses"][key] = clean_number(result["expenses"][key])
            if "expense_breakdown" in result["expenses"]:
                for item in result["expenses"]["expense_breakdown"]:
                    if "amount" in item:
                        item["amount"] = clean_number(item["amount"])
                    if "percentage" in item:
                        item["percentage"] = clean_number(item["percentage"])
        # Clean profitability numbers
        if "profitability" in result:
            for key in ["gross_profit", "operating_profit", "profit_before_tax", "net_profit", "gross_margin", "net_margin"]:
                if key in result["profitability"]:
                    result["profitability"][key] = clean_number(result["profitability"][key])
        
        # Compute dashboard metrics
        total_revenue = result.get("revenue", {}).get("total_revenue", 0) or 0
        total_expenses = result.get("expenses", {}).get("total_expenses", 0) or 0
        gross_profit = result.get("profitability", {}).get("gross_profit", 0) or 0
        net_profit = result.get("profitability", {}).get("net_profit", 0) or 0
        expense_breakdown = result.get("expenses", {}).get("expense_breakdown", [])
        
        # Top 5 expenses
        top_expenses = sorted(expense_breakdown, key=lambda x: x.get("amount", 0) or 0, reverse=True)[:5]
        
        # Calculate ratios
        gross_margin = result.get("profitability", {}).get("gross_margin", 0) or 0
        net_margin = result.get("profitability", {}).get("net_margin", 0) or 0
        expense_to_revenue = (total_expenses / total_revenue * 100) if total_revenue > 0 else 0
        
        result["dashboard"] = {
            "tables": {
                "expense_breakdown": expense_breakdown,
                "top_5_expenses": top_expenses
            },
            "charts": {
                "revenue_expense_trend": [{"period": result.get("period"), "revenue": total_revenue, "expenses": total_expenses}],
                "profit_margins": {
                    "gross_profit": gross_profit,
                    "net_profit": net_profit,
                    "gross_margin": gross_margin,
                    "net_margin": net_margin
                },
                "expense_pie": [{"category": e.get("category"), "amount": e.get("amount", 0) or 0} for e in expense_breakdown]
            },
            "ratios": {
                "gross_margin": round(gross_margin, 2),
                "net_margin": round(net_margin, 2),
                "expense_to_revenue_ratio": round(expense_to_revenue, 2)
            },
            "requires_multiple_pdfs": True
        }
        
        # Generate all required reports for Profit & Loss
        try:
            reports = generate_profit_loss_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting profit & loss: {str(e)}")

def extract_invoice(text):
    """Extract Invoice data with person-relevant details"""
    prompt = f"""Extract Invoice data from the following OCR text.

Return STRICT JSON:
{{
  "invoice_number": "",
  "invoice_date": "",
  "due_date": "",
  "seller": {{
    "name": "",
    "address": "",
    "gstin": "",
    "pan": "",
    "contact": "",
    "bank_details": ""
  }},
  "buyer": {{
    "name": "",
    "address": "",
    "gstin": "",
    "pan": "",
    "contact": ""
  }},
  "amounts": {{
    "subtotal": "",
    "discount": "",
    "taxable_amount": "",
    "cgst": "",
    "sgst": "",
    "igst": "",
    "cess": "",
    "tax_amount": "",
    "round_off": "",
    "total_amount": ""
  }},
  "payment": {{
    "payment_terms": "",
    "payment_mode": "",
    "payment_status": "",
    "paid_amount": "",
    "outstanding_amount": ""
  }},
  "line_items": [
    {{
      "item_code": "",
      "description": "",
      "hsn_sac": "",
      "quantity": "",
      "unit": "",
      "rate": "",
      "discount": "",
      "amount": "",
      "tax_rate": "",
      "tax_amount": ""
    }}
  ]
}}

Rules:
- Clean all numbers
- Dates in YYYY-MM-DD format
- Extract complete seller and buyer details
- Extract all tax components
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        # Clean amounts
        if "amounts" in result:
            for key in ["subtotal", "discount", "taxable_amount", "cgst", "sgst", "igst", "cess", "tax_amount", "round_off", "total_amount"]:
                if key in result["amounts"]:
                    result["amounts"][key] = clean_number(result["amounts"][key])
        # Clean payment amounts
        if "payment" in result:
            for key in ["paid_amount", "outstanding_amount"]:
                if key in result["payment"]:
                    result["payment"][key] = clean_number(result["payment"][key])
        # Clean line items
        if "line_items" in result:
            for item in result["line_items"]:
                for key in ["quantity", "rate", "discount", "amount", "tax_rate", "tax_amount"]:
                    if key in item:
                        item[key] = clean_number(item[key])
        
        # Compute dashboard metrics
        total_amount = result.get("amounts", {}).get("total_amount", 0) or 0
        line_items = result.get("line_items", [])
        cgst = result.get("amounts", {}).get("cgst", 0) or 0
        sgst = result.get("amounts", {}).get("sgst", 0) or 0
        igst = result.get("amounts", {}).get("igst", 0) or 0
        
        # Calculate average order value
        avg_order_value = total_amount / len(line_items) if line_items else 0
        
        # Tax breakdown
        tax_breakdown = []
        if cgst > 0:
            tax_breakdown.append({"type": "CGST", "amount": cgst})
        if sgst > 0:
            tax_breakdown.append({"type": "SGST", "amount": sgst})
        if igst > 0:
            tax_breakdown.append({"type": "IGST", "amount": igst})
        
        result["dashboard"] = {
            "tables": {
                "line_items": line_items
            },
            "charts": {
                "tax_breakdown": tax_breakdown,
                "buyer_seller_summary": {
                    "seller": result.get("seller", {}).get("name", ""),
                    "buyer": result.get("buyer", {}).get("name", ""),
                    "amount": total_amount
                }
            },
            "ratios": {
                "avg_order_value": round(avg_order_value, 2),
                "tax_rate_consistency": round((result.get("amounts", {}).get("tax_amount", 0) or 0) / (result.get("amounts", {}).get("taxable_amount", 0) or 1) * 100, 2) if result.get("amounts", {}).get("taxable_amount", 0) else 0
            },
            "requires_multiple_pdfs": True
        }
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting invoice: {str(e)}")

def extract_purchase_order(text):
    """Extract Purchase Order data with person-relevant details"""
    prompt = f"""Extract Purchase Order data from the following OCR text.

Return STRICT JSON:
{{
  "po_number": "",
  "po_date": "",
  "delivery_date": "",
  "validity_date": "",
  "vendor": {{
    "name": "",
    "address": "",
    "contact": "",
    "gstin": "",
    "pan": ""
  }},
  "buyer": {{
    "name": "",
    "address": "",
    "contact": "",
    "gstin": "",
    "pan": ""
  }},
  "items": [
    {{
      "item_code": "",
      "description": "",
      "specification": "",
      "quantity": "",
      "unit": "",
      "rate": "",
      "discount": "",
      "amount": "",
      "tax_rate": "",
      "tax_amount": ""
    }}
  ],
  "amounts": {{
    "subtotal": "",
    "discount": "",
    "taxable_amount": "",
    "cgst": "",
    "sgst": "",
    "igst": "",
    "tax_amount": "",
    "shipping_charges": "",
    "total_amount": ""
  }},
  "terms": {{
    "payment_terms": "",
    "delivery_terms": "",
    "warranty": "",
    "notes": ""
  }}
}}

Rules:
- Clean all numbers
- Dates in YYYY-MM-DD format
- Extract complete vendor and buyer details
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        # Clean amounts
        if "amounts" in result:
            for key in ["subtotal", "discount", "taxable_amount", "cgst", "sgst", "igst", "tax_amount", "shipping_charges", "total_amount"]:
                if key in result["amounts"]:
                    result["amounts"][key] = clean_number(result["amounts"][key])
        # Clean line items
        if "items" in result:
            for item in result["items"]:
                for key in ["quantity", "rate", "discount", "amount", "tax_rate", "tax_amount"]:
                    if key in item:
                        item[key] = clean_number(item[key])
        
        # Compute dashboard metrics
        total_amount = result.get("amounts", {}).get("total_amount", 0) or 0
        items = result.get("items", [])
        vendor_name = result.get("vendor", {}).get("name", "")
        
        # Item quantity summary
        item_quantities = [{"description": item.get("description"), "quantity": item.get("quantity", 0) or 0} for item in items]
        
        result["dashboard"] = {
            "tables": {
                "items": items
            },
            "charts": {
                "vendor_spend": [{"vendor": vendor_name, "amount": total_amount}],
                "item_quantities": item_quantities
            },
            "ratios": {
                "vendor_concentration": 100.0,  # Single PO = 100%
                "purchase_growth": 0  # Requires multiple POs
            },
            "requires_multiple_pdfs": False
        }
        
        # Generate all required reports for Purchase Order
        try:
            reports = generate_purchase_order_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting purchase order: {str(e)}")

def extract_salary_slip(text):
    """Extract Salary Slip data with person-relevant details"""
    prompt = f"""Extract Salary Slip/Payslip data from the following OCR text.

Return STRICT JSON:
{{
  "employee": {{
    "name": "",
    "employee_id": "",
    "designation": "",
    "department": "",
    "uan": "",
    "pan": "",
    "aadhaar": ""
  }},
  "employer": {{
    "name": "",
    "address": "",
    "pf_number": "",
    "esi_number": ""
  }},
  "salary_period": {{
    "month": "",
    "year": "",
    "days_paid": "",
    "working_days": ""
  }},
  "earnings": {{
    "basic_pay": "",
    "hra": "",
    "conveyance_allowance": "",
    "medical_allowance": "",
    "special_allowance": "",
    "bonus": "",
    "overtime": "",
    "other_allowances": "",
    "total_earnings": "",
    "gross_salary": ""
  }},
  "deductions": {{
    "provident_fund": "",
    "professional_tax": "",
    "income_tax": "",
    "esi": "",
    "loan_deduction": "",
    "other_deductions": "",
    "total_deductions": ""
  }},
  "net_salary": "",
  "year_to_date": {{
    "gross_ytd": "",
    "deductions_ytd": "",
    "net_ytd": ""
  }}
}}

Rules:
- Clean all numbers
- Dates in YYYY-MM-DD format
- Extract all earnings and deductions separately
- Include YTD (Year to Date) figures if available
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        # Clean earnings
        if "earnings" in result:
            for key in ["basic_pay", "hra", "conveyance_allowance", "medical_allowance", "special_allowance", "bonus", "overtime", "other_allowances", "total_earnings", "gross_salary"]:
                if key in result["earnings"]:
                    result["earnings"][key] = clean_number(result["earnings"][key])
        # Clean deductions
        if "deductions" in result:
            for key in ["provident_fund", "professional_tax", "income_tax", "esi", "loan_deduction", "other_deductions", "total_deductions"]:
                if key in result["deductions"]:
                    result["deductions"][key] = clean_number(result["deductions"][key])
        # Clean net salary and YTD
        if "net_salary" in result:
            result["net_salary"] = clean_number(result["net_salary"])
        if "year_to_date" in result:
            for key in ["gross_ytd", "deductions_ytd", "net_ytd"]:
                if key in result["year_to_date"]:
                    result["year_to_date"][key] = clean_number(result["year_to_date"][key])
        
        # Compute dashboard metrics
        gross_salary = result.get("earnings", {}).get("gross_salary", 0) or 0
        total_deductions = result.get("deductions", {}).get("total_deductions", 0) or 0
        basic_pay = result.get("earnings", {}).get("basic_pay", 0) or 0
        allowances = (result.get("earnings", {}).get("hra", 0) or 0) + (result.get("earnings", {}).get("conveyance_allowance", 0) or 0) + (result.get("earnings", {}).get("medical_allowance", 0) or 0) + (result.get("earnings", {}).get("special_allowance", 0) or 0)
        
        # Earnings breakdown
        earnings_breakdown = [
            {"type": "Basic Pay", "amount": basic_pay},
            {"type": "HRA", "amount": result.get("earnings", {}).get("hra", 0) or 0},
            {"type": "Allowances", "amount": allowances},
            {"type": "Bonus", "amount": result.get("earnings", {}).get("bonus", 0) or 0}
        ]
        
        # Deductions breakdown
        deductions_breakdown = [
            {"type": "Provident Fund", "amount": result.get("deductions", {}).get("provident_fund", 0) or 0},
            {"type": "Income Tax", "amount": result.get("deductions", {}).get("income_tax", 0) or 0},
            {"type": "Professional Tax", "amount": result.get("deductions", {}).get("professional_tax", 0) or 0},
            {"type": "ESI", "amount": result.get("deductions", {}).get("esi", 0) or 0}
        ]
        
        # Calculate ratios
        deduction_ratio = (total_deductions / gross_salary * 100) if gross_salary > 0 else 0
        allowance_to_basic = (allowances / basic_pay * 100) if basic_pay > 0 else 0
        ctc_monthly = gross_salary
        ctc_annual = gross_salary * 12
        
        result["dashboard"] = {
            "tables": {
                "earnings": earnings_breakdown,
                "deductions": deductions_breakdown
            },
            "charts": {
                "salary_trend": [{"month": result.get("salary_period", {}).get("month", ""), "net_salary": result.get("net_salary", 0) or 0}],
                "earnings_breakdown": earnings_breakdown,
                "deductions_breakdown": deductions_breakdown
            },
            "ratios": {
                "deduction_ratio": round(deduction_ratio, 2),
                "allowance_to_basic_ratio": round(allowance_to_basic, 2),
                "ctc_monthly": round(ctc_monthly, 2),
                "ctc_annual": round(ctc_annual, 2)
            },
            "requires_multiple_pdfs": True
        }
        
        # Generate all required reports for Salary Slip
        try:
            reports = generate_salary_slip_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting salary slip: {str(e)}")

def extract_balance_sheet(text):
    """Extract Balance Sheet data with person-relevant details"""
    prompt = f"""Extract Balance Sheet data from the following OCR text.

CRITICAL REQUIREMENTS:
1. Extract ONLY values that are explicitly present in the document
2. Do NOT use placeholder values like 0, "N/A", or any default numbers unless the document shows zero
3. If a value is missing in the document, use null (not 0, not empty string)
4. All numbers must be actual extracted values - clean them (remove currency symbols, commas) but do not modify the actual numeric value
5. Verify balance sheet equation: Total Assets MUST equal Total Liabilities + Total Equity (within rounding tolerance)
6. If breakdown items are provided, verify that subcategory amounts sum to category totals
7. Extract detailed breakdowns from the document if available (e.g., individual asset items, specific liability items)

Return STRICT JSON:
{{
  "entity_name": "",
  "entity_type": "",
  "period": "",
  "prepared_date": "",
  "assets": {{
    "current_assets": number or null,
    "fixed_assets": number or null,
    "intangible_assets": number or null,
    "investments": number or null,
    "other_assets": number or null,
    "total_assets": number or null,
    "asset_breakdown": [
      {{
        "category": "",
        "subcategory": "",
        "amount": number or null
      }}
    ]
  }},
  "liabilities": {{
    "current_liabilities": number or null,
    "long_term_liabilities": number or null,
    "loans": number or null,
    "creditors": number or null,
    "other_liabilities": number or null,
    "total_liabilities": number or null,
    "liability_breakdown": [
      {{
        "category": "",
        "subcategory": "",
        "amount": number or null
      }}
    ]
  }},
  "equity": {{
    "share_capital": number or null,
    "reserves": number or null,
    "retained_earnings": number or null,
    "total_equity": number or null
  }},
  "summary": {{
    "total_liabilities_and_equity": number or null,
    "is_balanced": true or false
  }}
}}

Rules:
- Clean all numbers: remove ₹, commas, but preserve actual numeric value (e.g., "₹2,50,000" -> 250000)
- Categorize assets accurately: Current, Fixed, Intangible, Investments based on what's in the document
- Categorize liabilities accurately: Current, Long-term, Loans based on what's in the document
- Verify calculations: total_assets should equal sum of all asset categories, total_liabilities should equal sum of all liability categories
- Verify balance: total_assets should equal total_liabilities + total_equity (set is_balanced = true only if this matches)
- Extract detailed breakdowns if available in document (individual line items)
- If a field is not found in document, use null (not 0, not empty string)
- Return ONLY valid JSON, no explanations or additional text

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        # Clean assets
        if "assets" in result:
            for key in ["current_assets", "fixed_assets", "intangible_assets", "investments", "other_assets", "total_assets"]:
                if key in result["assets"]:
                    result["assets"][key] = clean_number(result["assets"][key])
            if "asset_breakdown" in result["assets"]:
                for item in result["assets"]["asset_breakdown"]:
                    if "amount" in item:
                        item["amount"] = clean_number(item["amount"])
        # Clean liabilities
        if "liabilities" in result:
            for key in ["current_liabilities", "long_term_liabilities", "loans", "creditors", "other_liabilities", "total_liabilities"]:
                if key in result["liabilities"]:
                    result["liabilities"][key] = clean_number(result["liabilities"][key])
            if "liability_breakdown" in result["liabilities"]:
                for item in result["liabilities"]["liability_breakdown"]:
                    if "amount" in item:
                        item["amount"] = clean_number(item["amount"])
        # Clean equity
        if "equity" in result:
            for key in ["share_capital", "reserves", "retained_earnings", "total_equity"]:
                if key in result["equity"]:
                    result["equity"][key] = clean_number(result["equity"][key])
        # Clean summary
        if "summary" in result:
            if "total_liabilities_and_equity" in result["summary"]:
                result["summary"]["total_liabilities_and_equity"] = clean_number(result["summary"]["total_liabilities_and_equity"])
        
        # Compute dashboard metrics
        current_assets = result.get("assets", {}).get("current_assets", 0) or 0
        current_liabilities = result.get("liabilities", {}).get("current_liabilities", 0) or 0
        total_assets = result.get("assets", {}).get("total_assets", 0) or 0
        total_liabilities = result.get("liabilities", {}).get("total_liabilities", 0) or 0
        total_equity = result.get("equity", {}).get("total_equity", 0) or 0
        loans = result.get("liabilities", {}).get("loans", 0) or 0
        
        # Calculate ratios
        current_ratio = (current_assets / current_liabilities) if current_liabilities > 0 else 0
        debt_to_equity = (loans / total_equity) if total_equity > 0 else 0
        working_capital = current_assets - current_liabilities
        
        # Asset breakdown for pie chart
        asset_pie = [
            {"type": "Current Assets", "amount": current_assets},
            {"type": "Fixed Assets", "amount": result.get("assets", {}).get("fixed_assets", 0) or 0},
            {"type": "Investments", "amount": result.get("assets", {}).get("investments", 0) or 0},
            {"type": "Other Assets", "amount": result.get("assets", {}).get("other_assets", 0) or 0}
        ]
        
        # Debt vs Equity
        debt_equity_chart = {
            "debt": loans,
            "equity": total_equity
        }
        
        result["dashboard"] = {
            "tables": {
                "asset_breakdown": result.get("assets", {}).get("asset_breakdown", []),
                "liability_breakdown": result.get("liabilities", {}).get("liability_breakdown", [])
            },
            "charts": {
                "assets_vs_liabilities": {
                    "total_assets": total_assets,
                    "total_liabilities": total_liabilities,
                    "total_equity": total_equity
                },
                "debt_equity": debt_equity_chart,
                "asset_pie": asset_pie
            },
            "ratios": {
                "current_ratio": round(current_ratio, 2),
                "debt_to_equity": round(debt_to_equity, 2),
                "working_capital": round(working_capital, 2),
                "quick_ratio": round(current_ratio, 2),  # Simplified, would need inventory data
                "asset_turnover": 0  # Requires revenue data
            },
            "requires_multiple_pdfs": False
        }
        
        # Generate all required reports for Balance Sheet
        try:
            reports = generate_balance_sheet_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting balance sheet: {str(e)}")

def extract_audit_papers(text):
    """Extract Audit Papers data"""
    prompt = f"""Extract Audit Papers data from the following OCR text.

Return STRICT JSON:
{{
  "audit_period": "",
  "entity_name": "",
  "audit_type": "",
  "auditor_name": "",
  "key_findings": [],
  "adjustments": [],
  "recommendations": [],
  "audit_opinion": ""
}}

Rules:
- Extract all audit-related information
- Identify key findings and recommendations
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        
        # Generate all required reports for Audit Papers
        try:
            reports = generate_audit_papers_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting audit papers: {str(e)}")

def extract_agreement_contract(text):
    """Extract Agreement/Contract data"""
    prompt = f"""Extract Agreement/Contract data from the following OCR text.

Return STRICT JSON:
{{
  "contract_type": "",
  "parties": [
    {{"name": "", "role": "", "address": ""}}
  ],
  "contract_date": "",
  "validity_period": "",
  "key_terms": [],
  "obligations": [],
  "risks": [],
  "compliance_requirements": [],
  "summary": ""
}}

Rules:
- Extract all contract details
- Identify parties, terms, obligations, risks
- Return ONLY JSON

OCR Text:
{text[:8000]}
"""
    try:
        response_text = generate_content_with_vertexai(prompt, require_json=True)
        result = json.loads(response_text)
        
        # Generate all required reports for Agreement/Contract
        try:
            reports = generate_agreement_contract_reports(result, text)
            result["reports"] = reports
        except Exception as e:
            result["reports"] = {"error": f"Could not generate reports: {str(e)}"}
        
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting agreement/contract: {str(e)}")

# ------------------------------
# MAIN PROCESSING ROUTE
# ------------------------------

@app.post("/process-audit")
async def process_audit_files(
    files: List[UploadFile] = File(...),
    document_type: Optional[str] = Form("audit")
):
    """Process multiple PDF files for comprehensive audit report generation
    
    File validation:
    - Max file size per file: 50MB
    - Allowed file types: PDF, DOCX, DOC, XLSX, XLS, JPG, JPEG, PNG
    """
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.jpg', '.jpeg', '.png'}
    
    print(f"\n{'='*60}")
    print(f"POST /process-audit - Processing {len(files) if files else 0} files for audit report")
    print(f"{'='*60}\n")
    
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="At least one file is required for audit report")
    
    temp_files = {}
    extracted_texts = {}
    
    try:
        import tempfile
        temp_dir = tempfile.gettempdir()
        
        # Process each file
        for file in files:
            if not file.filename:
                continue
            
            # Validate file extension
            file_extension = os.path.splitext(file.filename)[1].lower()
            if file_extension not in ALLOWED_EXTENSIONS:
                allowed_types_str = ', '.join(sorted([ext.upper() for ext in ALLOWED_EXTENSIONS]))
                raise HTTPException(
                    status_code=400,
                    detail=f"File '{file.filename}' has unsupported type '{file_extension}'. Allowed types: {allowed_types_str}"
                )
                
            print(f"Processing file: {file.filename}")
            # Sanitize filename for filesystem
            safe_filename = file.filename.replace("/", "_").replace("\\", "_")
            temp_file_path = os.path.join(temp_dir, f"finsight_audit_{os.getpid()}_{safe_filename}")
            
            # Read and validate file size
            content = b""
            total_size = 0
            chunk_size = 1024 * 1024  # 1MB chunks
            
            while True:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=413,
                        detail=f"File '{file.filename}' size ({total_size / (1024*1024):.2f}MB) exceeds maximum allowed size of 50MB."
                    )
                content += chunk
            
            if total_size == 0:
                raise HTTPException(status_code=400, detail=f"File '{file.filename}' is empty.")
            
            print(f"✓ File validated: {file.filename} ({total_size / (1024*1024):.2f}MB)")
            print(f"Read {len(content)} bytes from {file.filename}")
            with open(temp_file_path, "wb") as temp:
                temp.write(content)
            temp_files[file.filename] = temp_file_path
            print(f"Saved to: {temp_file_path}")
            
            # Extract text
            print(f"Extracting text from {file.filename}...")
            try:
                if file.filename.lower().endswith(".pdf"):
                    text = extract_text_from_pdf(temp_file_path)
                elif file.filename.lower().endswith((".docx", ".doc")):
                    text = extract_text_from_docx(temp_file_path)
                elif file.filename.lower().endswith((".jpg", ".jpeg", ".png")):
                    text = extract_text_from_image(temp_file_path)
                else:
                    # Try to read as text file
                    with open(temp_file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
            except Exception as extract_error:
                print(f"Error extracting text from {file.filename}: {str(extract_error)}")
                text = f"Error extracting text: {str(extract_error)}"
            
            # Map filename to document type based on field name or filename
            filename_lower = file.filename.lower()
            doc_type = "unknown"
            
            # Check field name if available (from form data)
            if hasattr(file, 'field_name'):
                field_name = file.field_name.lower()
                if "trial" in field_name or "balance" in field_name:
                    doc_type = "trial_balance"
                elif "profit" in field_name or "loss" in field_name:
                    doc_type = "profit_loss"
                elif "balance" in field_name and "sheet" in field_name:
                    doc_type = "balance_sheet"
                elif "ledger" in field_name:
                    doc_type = "general_ledger"
                elif "cash" in field_name:
                    doc_type = "cash_book"
                elif "bank" in field_name:
                    doc_type = "bank_statement"
                elif "asset" in field_name or "fixed" in field_name:
                    doc_type = "fixed_asset_register"
                elif "gst" in field_name:
                    doc_type = "gst_returns"
                elif "tds" in field_name:
                    doc_type = "tds_summary"
            
            # Fallback to filename analysis
            if doc_type == "unknown":
                if "trial" in filename_lower or ("balance" in filename_lower and "sheet" not in filename_lower):
                    doc_type = "trial_balance"
                elif "profit" in filename_lower or "loss" in filename_lower or "p&l" in filename_lower or "p_l" in filename_lower:
                    doc_type = "profit_loss"
                elif "balance" in filename_lower and "sheet" in filename_lower:
                    doc_type = "balance_sheet"
                elif "ledger" in filename_lower:
                    doc_type = "general_ledger"
                elif "cash" in filename_lower:
                    doc_type = "cash_book"
                elif "bank" in filename_lower:
                    doc_type = "bank_statement"
                elif "asset" in filename_lower or "fixed" in filename_lower:
                    doc_type = "fixed_asset_register"
                elif "gst" in filename_lower:
                    doc_type = "gst_returns"
                elif "tds" in filename_lower:
                    doc_type = "tds_summary"
            
            extracted_texts[doc_type] = {
                "filename": file.filename,
                "text": text,
                "type": doc_type
            }
            print(f"Extracted {len(text)} characters from {file.filename} (type: {doc_type})")
        
        print(f"\nExtracted texts from {len(extracted_texts)} documents")
        print(f"Document types: {list(extracted_texts.keys())}")
        
        # Generate comprehensive audit report
        print("\nGenerating comprehensive audit report...")
        result = generate_comprehensive_audit_report(extracted_texts)
        
        print(f"\nAudit report generated successfully!")
        print(f"Report keys: {list(result.keys()) if isinstance(result, dict) else 'Not a dict'}")
        
        return JSONResponse(content=result)
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"\n{'='*60}")
        print("ERROR in process_audit_files:")
        print(f"{'='*60}")
        traceback.print_exc()
        print(f"{'='*60}\n")
        raise HTTPException(status_code=500, detail=f"Error processing audit files: {str(e)}")
    
    finally:
        # Clean up temporary files
        for temp_path in temp_files.values():
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                    print(f"Cleaned up: {temp_path}")
                except Exception as cleanup_error:
                    print(f"Warning: Could not delete {temp_path}: {cleanup_error}")

@app.post("/process-gst")
async def process_gst_files(
    files: List[UploadFile] = File(...),
    document_type: Optional[str] = Form("gst_return")
):
    """Process multiple Excel files for GST report generation (GSTR-2B, Purchase Register, Vendor Master)
    
    File validation:
    - Max file size per file: 50MB
    - Only Excel files allowed: .xlsx, .xls
    """
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_EXTENSIONS = {'.xlsx', '.xls'}
    
    print(f"\n{'='*60}")
    print(f"POST /process-gst - Processing {len(files) if files else 0} Excel files for GST reports")
    print(f"{'='*60}\n")
    
    if not files or len(files) == 0:
        raise HTTPException(status_code=400, detail="At least one Excel file is required for GST reports")
    
    temp_files = {}
    excel_data_dict = {}
    
    try:
        import tempfile
        temp_dir = tempfile.gettempdir()
        
        # Process each Excel file
        for file in files:
            if not file.filename:
                continue
            
            # Validate file extension (only Excel files)
            file_extension = os.path.splitext(file.filename)[1].lower()
            if file_extension not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"File '{file.filename}' is not an Excel file. Only .xlsx and .xls files are allowed for GST processing."
                )
            
            print(f"Processing Excel file: {file.filename}")
            safe_filename = file.filename.replace("/", "_").replace("\\", "_")
            temp_file_path = os.path.join(temp_dir, f"finsight_gst_{os.getpid()}_{safe_filename}")
            
            # Read and validate file size
            content = b""
            total_size = 0
            chunk_size = 1024 * 1024  # 1MB chunks
            
            while True:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=413,
                        detail=f"File '{file.filename}' size ({total_size / (1024*1024):.2f}MB) exceeds maximum allowed size of 50MB."
                    )
                content += chunk
            
            if total_size == 0:
                raise HTTPException(status_code=400, detail=f"File '{file.filename}' is empty.")
            
            print(f"✓ File validated: {file.filename} ({total_size / (1024*1024):.2f}MB)")
            print(f"Read {len(content)} bytes from {file.filename}")
            with open(temp_file_path, "wb") as temp:
                temp.write(content)
            temp_files[file.filename] = temp_file_path
            print(f"Saved to: {temp_file_path}")
            
            # Extract data from Excel
            print(f"Extracting data from {file.filename}...")
            try:
                excel_data = extract_data_from_excel(temp_file_path)
                
                # Map filename to file type
                filename_lower = file.filename.lower()
                file_type = "unknown"
                
                if "gstr" in filename_lower or "2b" in filename_lower or "summary" in filename_lower:
                    file_type = "gstr2b_summary"
                elif "purchase" in filename_lower or "register" in filename_lower:
                    file_type = "purchase_register"
                elif "vendor" in filename_lower or "master" in filename_lower:
                    file_type = "vendor_master"
                
                excel_data_dict[file_type] = excel_data
                print(f"Extracted data from {file.filename} (type: {file_type})")
                print(f"  Sheets: {excel_data.get('sheet_names', [])}")
                
                # Log sample of actual data extracted
                excel_data_obj = excel_data.get("excel_data", {})
                for sheet_name, sheet_info in excel_data_obj.items():
                    row_count = sheet_info.get("row_count", 0)
                    columns = sheet_info.get("columns", [])
                    print(f"  Sheet '{sheet_name}': {row_count} rows, Columns: {columns}")
                    if sheet_info.get("data"):
                        print(f"  First row sample: {sheet_info['data'][0] if len(sheet_info['data']) > 0 else 'No data'}")
                
            except Exception as extract_error:
                print(f"Error extracting data from {file.filename}: {str(extract_error)}")
                import traceback
                traceback.print_exc()
                excel_data_dict[file.filename] = {
                    "error": f"Error extracting data: {str(extract_error)}",
                    "summary_text": f"Error reading {file.filename}",
                    "sheet_names": []
                }
        
        print(f"\nExtracted data from {len(excel_data_dict)} Excel files")
        print(f"File types: {list(excel_data_dict.keys())}")
        
        # Generate comprehensive GST report from Excel data
        print("\nGenerating comprehensive GST report from Excel data...")
        gst_report = generate_gst_reports_from_excel(excel_data_dict)
        
        print(f"\nGST report generated successfully!")
        print(f"Report keys: {list(gst_report.keys()) if isinstance(gst_report, dict) else 'Not a dict'}")
        
        return JSONResponse(content=gst_report)
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"\n{'='*60}")
        print("ERROR in process_gst_files:")
        print(f"{'='*60}")
        traceback.print_exc()
        print(f"{'='*60}\n")
        raise HTTPException(status_code=500, detail=f"Error processing GST Excel files: {str(e)}")
    
    finally:
        # Clean up temporary files
        for temp_path in temp_files.values():
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                    print(f"Cleaned up: {temp_path}")
                except Exception as cleanup_error:
                    print(f"Warning: Could not delete {temp_path}: {cleanup_error}")

@app.post("/process")
async def process_file(
    file: UploadFile = File(...),
    document_type: Optional[str] = Form(None)
):
    """
    Process document synchronously and return results immediately.
    Uses Redis cache to return cached results instantly for identical documents.
    
    File validation:
    - Max file size: 50MB
    - Allowed file types: PDF, DOCX, DOC, XLSX, XLS, JPG, JPEG, PNG
    """
    # File validation constants
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB in bytes
    ALLOWED_EXTENSIONS = {
        '.pdf', '.docx', '.doc', 
        '.xlsx', '.xls', 
        '.jpg', '.jpeg', '.png'
    }
    
    # Validate filename
    filename = file.filename
    if not filename:
        raise HTTPException(
            status_code=400,
            detail="No filename provided. Please ensure the file has a valid name."
        )
    
    # Validate file extension
    file_extension = os.path.splitext(filename)[1].lower()
    if file_extension not in ALLOWED_EXTENSIONS:
        allowed_types_str = ', '.join(sorted([ext.upper() for ext in ALLOWED_EXTENSIONS]))
        raise HTTPException(
            status_code=400,
            detail=f"File type '{file_extension}' is not supported. Allowed file types: {allowed_types_str}. Please upload a supported file format."
        )
    
    print(f"\n{'='*60}")
    print(f"POST /process - Processing file: {filename}")
    print(f"File extension: {file_extension}")
    print(f"Document type: {document_type}")
    print(f"{'='*60}\n")
    
    temp_file_path = None
    
    try:
        # Read file content with size validation
        content = b""
        total_size = 0
        chunk_size = 1024 * 1024  # Read in 1MB chunks
        
        while True:
            chunk = await file.read(chunk_size)
            if not chunk:
                break
            total_size += len(chunk)
            
            # Check size during read to avoid loading huge files into memory
            if total_size > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size ({total_size / (1024*1024):.2f}MB) exceeds maximum allowed size of 50MB. Please upload a smaller file."
                )
            
            content += chunk
        
        # Final size validation
        if total_size == 0:
            raise HTTPException(
                status_code=400,
                detail="File is empty. Please upload a file with content."
            )
        
        if total_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File size ({total_size / (1024*1024):.2f}MB) exceeds maximum allowed size of 50MB. Please upload a smaller file."
            )
        
        print(f"✓ File validation passed: {filename} ({total_size / (1024*1024):.2f}MB, {file_extension})")
        
        # Check cache for complete document processing result
        if REDIS_AVAILABLE and redis_client:
            try:
                doc_cache_key = get_document_cache_key(content, document_type)
                cached_result = redis_client.get(doc_cache_key)
                if cached_result:
                    print(f"✓ CACHE HIT: Returning cached result for document (key: {doc_cache_key[:30]}...)")
                    print(f"  Document: {filename} | Type: {document_type or 'auto'}")
                    return JSONResponse(content=json.loads(cached_result))
                else:
                    print(f"✓ CACHE MISS: Processing new document (key: {doc_cache_key[:30]}...)")
            except Exception as e:
                print(f"Cache check error (continuing with processing): {str(e)}")
        
        # Save to temporary file
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_file_path = os.path.join(temp_dir, f"finsight_{filename}")
        
        with open(temp_file_path, "wb") as temp:
            temp.write(content)
        
        print(f"File saved to: {temp_file_path}")
        
        # Extract text
        print("Extracting text...")
        text = ""
        if filename.lower().endswith(".pdf"):
            text = extract_text_from_pdf(temp_file_path)
        elif filename.lower().endswith((".docx", ".doc")):
            text = extract_text_from_docx(temp_file_path)
        elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
            text = extract_text_from_image(temp_file_path)
        elif filename.lower().endswith((".xlsx", ".xls")):
            excel_data = extract_excel_data(temp_file_path)
            text = excel_data.get("summary_text", "")
        else:
            with open(temp_file_path, "r", encoding="utf-8") as f:
                text = f.read()
        
        print(f"Text extracted: {len(text)} characters")
        
        # Normalize document type
        if document_type:
            document_type = document_type.lower().replace(" ", "_").replace("-", "_")
        
        # Classify document if type not provided
        if not document_type:
            print("Classifying document type...")
            try:
                detected = classify_document(text)
                document_type = detected.get("type", "").lower().replace(" ", "_")
                print(f"Detected document type: {document_type}")
            except Exception as e:
                print(f"Warning: Classification failed: {str(e)}")
                document_type = "unknown"
        
        # Extract data based on document type
        print(f"Extracting data for type: {document_type}")
        result = {}
        
        if document_type == "bank_statement" or (not document_type and "bank" in text.lower()):
            result = extract_bank_statement_structured(text)
        elif document_type == "gst_return" or document_type == "gst_document":
            result = extract_gst_return(text)
        elif document_type == "trial_balance":
            result = extract_trial_balance(text)
        elif document_type == "profit_loss" or document_type == "p&l":
            result = extract_profit_loss(text)
        elif document_type == "invoice":
            result = extract_invoice(text)
        elif document_type == "purchase_order" or document_type == "po":
            result = extract_purchase_order(text)
        elif document_type == "salary_slip" or document_type == "payslip":
            result = extract_salary_slip(text)
        elif document_type == "balance_sheet":
            result = extract_balance_sheet(text)
        elif document_type == "audit_papers" or document_type == "audit":
            result = extract_audit_papers(text)
        elif document_type == "agreement_contract" or document_type in ["agreement", "contract"]:
            result = extract_agreement_contract(text)
        else:
            raise HTTPException(
                    status_code=400,
                detail=f"Unsupported document type: {document_type or 'unknown'}. Supported types: bank_statement, gst_return, trial_balance, profit_loss, invoice, purchase_order, salary_slip, balance_sheet, audit_papers, agreement_contract"
            )
        
        print(f"Processing completed successfully!")
        
        # Cache the complete result if Redis is available
        if REDIS_AVAILABLE and redis_client:
            try:
                doc_cache_key = get_document_cache_key(content, document_type)
                # Cache for 7 days (604800 seconds) - documents rarely change
                redis_client.setex(doc_cache_key, 604800, json.dumps(result))
                print(f"✓ Cached complete document result (key: {doc_cache_key[:30]}..., TTL: 7 days)")
            except Exception as e:
                print(f"Cache write error (result still returned): {str(e)}")
        
        return JSONResponse(content=result)
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Error processing file: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                print(f"Cleaned up temporary file: {temp_file_path}")
            except Exception as e:
                print(f"Warning: Could not delete temporary file {temp_file_path}: {e}")


@app.get("/job/{job_id}/status")
async def get_job_status(job_id: str):
    """Get the status of a processing job"""
    if not CELERY_AVAILABLE or not celery_app:
        raise HTTPException(
            status_code=503,
            detail="Job queue is not available. Please start Redis and Celery worker."
        )
    
    try:
        task = celery_app.AsyncResult(job_id)
        
        # Safely get task state - use ready() and successful() methods to avoid exceptions
        task_state = 'PENDING'
        task_ready = False
        task_info = {}
        
        try:
            # Check if task is ready (completed, failed, or revoked) - this doesn't throw exceptions
            task_ready = task.ready()
            
            if task_ready:
                # Task is done - check if successful or failed
                if task.successful():
                    task_state = 'SUCCESS'
                elif task.failed():
                    task_state = 'FAILURE'
                else:
                    # Try to get state, but don't fail if it throws
                    try:
                        task_state = task.state
                    except Exception:
                        task_state = 'UNKNOWN'
            else:
                # Task is still running or pending - check state safely
                try:
                    # For in-progress tasks, try to get state
                    task_state = task.state
                    if task_state not in ['PENDING', 'PROCESSING', 'STARTED']:
                        task_state = 'PENDING'
                except Exception:
                    task_state = 'PENDING'
        except Exception as state_error:
            print(f"ERROR getting task state for {job_id}: {str(state_error)}")
            # Default to pending if we can't determine state
            task_state = 'PENDING'
            task_ready = False
        
        # Safely get task info/meta - only access if task is ready or use safe methods
        if task_ready:
            try:
                # Use get() with propagate=False to safely retrieve result/info
                task_info_raw = task.get(propagate=False)
                if isinstance(task_info_raw, dict):
                    task_info = task_info_raw
                elif task_info_raw is not None:
                    task_info = {'result': task_info_raw}
            except Exception:
                try:
                    # Fallback: try to access info directly for failed tasks
                    if task_state == 'FAILURE':
                        info = task.info
                        if isinstance(info, dict):
                            task_info = info
                        elif info is not None:
                            task_info = {'error': str(info)}
                except Exception:
                    task_info = {}
        else:
            # For non-ready tasks, try to get progress updates
            try:
                info = task.info
                if isinstance(info, dict):
                    task_info = info
                elif info is not None and info != 'None':
                    task_info = {'info': str(info)}
            except Exception:
                task_info = {}
        
        if task_state == 'PENDING':
            # Job is waiting to start
            response = {
                'job_id': job_id,
                'status': 'pending',
                'state': task_state,
                'message': 'Job is waiting to be processed'
            }
        elif task_state == 'PROCESSING':
            # Job is currently being processed
            meta = task_info if isinstance(task_info, dict) else {}
            response = {
                'job_id': job_id,
                'status': 'processing',
                'state': task_state,
                'progress': meta.get('progress', 0),
                'message': meta.get('status', 'Processing document...')
            }
        elif task_state == 'SUCCESS':
            # Job completed successfully
            meta = task_info if isinstance(task_info, dict) else {}
            # Safely get result
            try:
                result = meta.get('result') if isinstance(meta, dict) else None
                if result is None:
                    result = task.result if hasattr(task, 'result') else None
            except Exception as result_error:
                print(f"WARNING: Error getting task result for {job_id}: {str(result_error)}")
                result = None
            
            response = {
                'job_id': job_id,
                'status': 'completed',
                'state': task_state,
                'progress': 100,
                'message': 'Processing completed successfully'
            }
            # Include result if available
            if result:
                response['result'] = result
        elif task_state == 'FAILURE':
            # Job failed
            meta = task_info if isinstance(task_info, dict) else {}
            error_message = "Unknown error"
            error_type = "UnknownError"
            
            if isinstance(meta, dict):
                error_message = meta.get('error', str(task_info) if task_info else "Unknown error")
                error_type = meta.get('error_type', 'UnknownError')
            elif task_info:
                error_message = str(task_info)
            
            response = {
                'job_id': job_id,
                'status': 'failed',
                'state': task_state,
                'error': error_message,
                'error_type': error_type,
                'message': f"Processing failed: {error_message}"
            }
        else:
            # Unknown state
            response = {
                'job_id': job_id,
                'status': 'unknown',
                'state': task_state,
                'message': f'Job is in {task_state} state'
            }
        
        return JSONResponse(content=response)
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_job_status for {job_id}: {str(e)}")
        print(f"Traceback:\n{error_trace}")
        raise HTTPException(
            status_code=500,
            detail=f"Error checking job status: {str(e)}. Check backend logs for details."
        )


@app.get("/job/{job_id}/result")
async def get_job_result(job_id: str):
    """Get the result of a completed processing job"""
    if not CELERY_AVAILABLE or not celery_app:
        raise HTTPException(
            status_code=503,
            detail="Job queue is not available. Please start Redis and Celery worker."
        )
    
    try:
        task = celery_app.AsyncResult(job_id)
        
        # Safely get task state
        try:
            task_state = task.state
        except Exception as state_error:
            print(f"ERROR getting task state for result {job_id}: {str(state_error)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error accessing task state: {str(state_error)}"
            )
        
        if task_state == 'PENDING':
            raise HTTPException(status_code=202, detail="Job is still pending")
        elif task_state == 'PROCESSING':
            raise HTTPException(status_code=202, detail="Job is still processing")
        elif task_state == 'FAILURE':
            try:
                meta = task.info if task.info else {}
                if isinstance(meta, dict):
                    error_msg = meta.get('error', 'Unknown error')
                else:
                    error_msg = str(meta) if meta else 'Unknown error'
            except Exception as e:
                error_msg = f"Error retrieving failure info: {str(e)}"
            raise HTTPException(
                    status_code=500,
                detail=f"Job failed: {error_msg}"
            )
        elif task_state == 'SUCCESS':
            # Safely get result - try task.result first (most reliable)
            result = None
            try:
                # Get result directly from task (this is the actual return value)
                result = task.result
                if result is None:
                    # Fallback: try getting from task.info meta
                    try:
                        task_info = task.info if task.info else {}
                        if isinstance(task_info, dict):
                            result = task_info.get('result')
                    except Exception:
                        pass
            except Exception as result_error:
                print(f"ERROR getting task result for {job_id}: {str(result_error)}")
                # Try fallback to meta
                try:
                    task_info = task.info if task.info else {}
                    if isinstance(task_info, dict):
                        result = task_info.get('result')
                except Exception:
                    pass
            
            if result is None:
                raise HTTPException(
                    status_code=404,
                    detail="Job completed but result is not available. The task may have completed but result was not properly stored."
                )
            
            return JSONResponse(content=result)
        else:
            raise HTTPException(status_code=500, detail=f"Unknown job state: {task_state}")
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_job_result for {job_id}: {str(e)}")
        print(f"Traceback:\n{error_trace}")
        raise HTTPException(status_code=500, detail=f"Error retrieving job result: {str(e)}")

@app.get("/health")
async def health():
    """Health check endpoint"""
    redis_status = "available" if REDIS_AVAILABLE else "unavailable"
    return {"status": "healthy", "redis": redis_status}

@app.get("/cache/stats")
async def cache_stats():
    """Get Redis cache statistics"""
    if not REDIS_AVAILABLE or not redis_client:
        raise HTTPException(status_code=503, detail="Redis is not available")
    
    try:
        # Get cache keys count for different cache types
        ai_cache_keys = redis_client.keys("gemini_cache:*")
        document_cache_keys = redis_client.keys("document_cache:*")
        ai_cache_count = len(ai_cache_keys)
        document_cache_count = len(document_cache_keys)
        total_cache_count = ai_cache_count + document_cache_count
        
        # Get Redis info
        info = redis_client.info("memory")
        memory_used = info.get("used_memory_human", "N/A")
        
        return {
            "status": "ok",
            "cache_entries": {
                "ai_responses": ai_cache_count,
                "document_results": document_cache_count,
                "total": total_cache_count
            },
            "memory_used": memory_used,
            "redis_available": True
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting cache stats: {str(e)}")

@app.post("/export/tally")
async def export_to_tally(request_data: Dict):
    """
    Export processed document data to Tally XML format
    
    Args:
        request_data: JSON body containing:
            - data: Processed document data (from /process endpoint)
            - document_type: Type of document (bank_statement, invoice, etc.)
            - company_name: Optional Tally company name (default: "FinSight Company")
    
    Returns:
        XML file ready for import into Tally ERP/TallyPrime
    """
    if not TALLY_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Tally integration is not available. Please check installation."
        )
    
    try:
        # Extract parameters from request body
        data = request_data.get("data")
        document_type = request_data.get("document_type")
        company_name = request_data.get("company_name")
        
        if not data:
            raise HTTPException(
                status_code=400,
                detail="Missing required field: 'data' in request body"
            )
        
        if not document_type:
            raise HTTPException(
                status_code=400,
                detail="Missing required field: 'document_type' in request body"
            )
        
        company = company_name or "FinSight Company"
        xml_content = export_to_tally_xml(data, document_type, company)
        
        # Create a temporary file for the XML
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.xml', delete=False, encoding='utf-8')
        temp_file.write(xml_content)
        temp_file.close()
        
        filename = f"tally_export_{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xml"
        return FileResponse(
            temp_file.name,
            media_type="application/xml",
            filename=filename,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Type": "application/xml; charset=utf-8"
            }
        )
    except Exception as e:
        import traceback
        print(f"Error exporting to Tally: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error generating Tally XML: {str(e)}"
        )


@app.post("/process/export-tally")
async def process_and_export_to_tally(
    file: UploadFile = File(...),
    document_type: Optional[str] = Form(None),
    company_name: Optional[str] = Form(None)
):
    """
    Process a document and export directly to Tally XML format
    
    This endpoint:
    1. Processes the uploaded document
    2. Converts the extracted data to Tally vouchers
    3. Returns XML file ready for import into Tally
    """
    if not TALLY_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Tally integration is not available."
        )
    
    # First, process the document using existing logic
    # We'll reuse the /process endpoint logic
    temp_file_path = None
    
    try:
        # Read and validate file
        content = await file.read()
        filename = file.filename
        
        # Validate file size and type (reuse validation logic)
        MAX_FILE_SIZE = 50 * 1024 * 1024
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File size exceeds maximum allowed size of 50MB"
            )
        
        # Save to temporary file
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_file_path = os.path.join(temp_dir, f"finsight_tally_{filename}")
        
        with open(temp_file_path, "wb") as temp:
            temp.write(content)
        
        # Extract text (reuse existing logic)
        text = ""
        if filename.lower().endswith(".pdf"):
            text = extract_text_from_pdf(temp_file_path)
        elif filename.lower().endswith((".docx", ".doc")):
            text = extract_text_from_docx(temp_file_path)
        elif filename.lower().endswith((".jpg", ".jpeg", ".png")):
            text = extract_text_from_image(temp_file_path)
        elif filename.lower().endswith((".xlsx", ".xls")):
            excel_data = extract_excel_data(temp_file_path)
            text = excel_data.get("summary_text", "")
        else:
            with open(temp_file_path, "r", encoding="utf-8") as f:
                text = f.read()
        
        # Normalize document type
        if document_type:
            document_type = document_type.lower().replace(" ", "_").replace("-", "_")
        
        # Extract data based on document type
        result = {}
        if document_type == "bank_statement":
            result = extract_bank_statement_structured(text)
        elif document_type == "invoice":
            result = extract_invoice(text)
        elif document_type == "trial_balance":
            result = extract_trial_balance(text)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Document type '{document_type}' is not supported for Tally export. Supported types: bank_statement, invoice, trial_balance"
            )
        
        # Export to Tally XML
        company = company_name or "FinSight Company"
        xml_content = export_to_tally_xml(result, document_type, company)
        
        # Create response file
        temp_xml_file = tempfile.NamedTemporaryFile(mode='w', suffix='.xml', delete=False, encoding='utf-8')
        temp_xml_file.write(xml_content)
        temp_xml_file.close()
        
        return FileResponse(
            temp_xml_file.name,
            media_type="application/xml",
            filename=f"tally_export_{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xml",
            headers={
                "Content-Disposition": f'attachment; filename="tally_export_{document_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xml"'
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Error in process_and_export_to_tally: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error processing and exporting to Tally: {str(e)}"
        )
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass


@app.get("/cache/clear")
async def clear_cache():
    """Clear all cached responses"""
    if not REDIS_AVAILABLE or not redis_client:
        raise HTTPException(status_code=503, detail="Redis is not available")
    
    try:
        ai_cache_keys = redis_client.keys("gemini_cache:*")
        document_cache_keys = redis_client.keys("document_cache:*")
        all_cache_keys = ai_cache_keys + document_cache_keys
        if all_cache_keys:
            redis_client.delete(*all_cache_keys)
        return {
            "status": "success",
            "cleared_entries": {
                "ai_responses": len(ai_cache_keys),
                "document_results": len(document_cache_keys),
                "total": len(all_cache_keys)
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error clearing cache: {str(e)}")

@app.get("/")
async def root():
    """Health check endpoint"""
    return {"status": "ok", "message": "FinSight Document Processor API"}

@app.get("/health")
async def health():
    """Health check endpoint"""
    return {"status": "healthy"}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

